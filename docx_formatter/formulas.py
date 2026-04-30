from __future__ import annotations

from lxml import etree

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .styles import ALIGNMENT_MAP
from .utils import parse_distance


MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def latex_to_omml_element(latex_str: str):
    """Convert LaTeX string to OMML XML element via MathML."""
    import latex2mathml.converter
    import mathml2omml

    mathml = latex2mathml.converter.convert(latex_str)
    omml_str = mathml2omml.convert(mathml)
    # mathml2omml outputs without namespace declaration, wrap it
    wrapped = f'<wrapper xmlns:m="{MATH_NS}">{omml_str}</wrapper>'
    root = etree.fromstring(wrapped)
    return root[0]  # return the m:oMath child


def set_formula_font(omml_element, font_en: str) -> None:
    """Set font on all m:r elements within an OMML element."""
    for mr in omml_element.iter(f"{{{MATH_NS}}}r"):
        existing_rpr = mr.find(qn("m:rPr"))
        if existing_rpr is not None:
            mr.remove(existing_rpr)
        rpr = etree.SubElement(mr, qn("m:rPr"))
        rfonts = etree.SubElement(rpr, qn("m:rFonts"))
        rfonts.set(qn("m:ascii"), font_en)
        rfonts.set(qn("m:hAnsi"), font_en)


def is_formula_paragraph(paragraph) -> bool:
    """Check if a paragraph contains m:oMath or m:oMathPara elements."""
    p = paragraph._p
    return (
        p.find(qn("m:oMath")) is not None
        or p.find(qn("m:oMathPara")) is not None
    )


def detect_formula_paragraphs(document) -> list[int]:
    """Return indices of paragraphs containing formulas."""
    return [
        i for i, p in enumerate(document.paragraphs) if is_formula_paragraph(p)
    ]


def insert_formula_paragraph(document, latex_str: str, font_en: str):
    """Add a new paragraph with formula from LaTeX, return the paragraph."""
    p = document.add_paragraph()
    omml = latex_to_omml_element(latex_str)
    set_formula_font(omml, font_en)
    p._p.append(omml)
    return p


def apply_formula_formatting(document, formula_config) -> None:
    """Apply formatting to all formula paragraphs."""
    for paragraph in document.paragraphs:
        if not is_formula_paragraph(paragraph):
            continue
        paragraph.paragraph_format.first_line_indent = Pt(0)
        # Don't set center alignment here — numbering uses tab stops instead
        before_unit, before_value = parse_distance(formula_config.space_before)
        after_unit, after_value = parse_distance(formula_config.space_after)
        if before_unit == "pt":
            paragraph.paragraph_format.space_before = Pt(before_value)
        if after_unit == "pt":
            paragraph.paragraph_format.space_after = Pt(after_value)


def _make_run_with_tab():
    """Create a w:r element containing a w:tab."""
    r = OxmlElement("w:r")
    r.append(OxmlElement("w:tab"))
    return r


def _set_tab_stops(paragraph, center_pos: int, right_pos: int) -> None:
    """Set center and right tab stops on a paragraph (positions in twips)."""
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:tabs"))
    if existing is not None:
        pPr.remove(existing)
    tabs = OxmlElement("w:tabs")
    center_tab = OxmlElement("w:tab")
    center_tab.set(qn("w:val"), "center")
    center_tab.set(qn("w:pos"), str(center_pos))
    tabs.append(center_tab)
    right_tab = OxmlElement("w:tab")
    right_tab.set(qn("w:val"), "right")
    right_tab.set(qn("w:pos"), str(right_pos))
    tabs.append(right_tab)
    pPr.append(tabs)


def apply_formula_numbering(document, formula_config, page_config=None, formula_items=None, registry=None) -> None:
    """Add chapter-based or continuous numbering to formula paragraphs.

    Uses center tab + right tab so formula is centered and number is right-aligned.

    Args:
        formula_items: list of (paragraph_index, item_dict) where item_dict has "key".
                       If None, all formula paragraphs are numbered without registration.
        registry: LabelRegistry for cross-references. If None, no registration occurs.
    """
    # Calculate tab stop positions from page config
    if page_config is not None:
        from .utils import cm_to_twips, text_area_twips
        left_cm = float(page_config.margins.left.removesuffix("cm"))
        right_cm = float(page_config.margins.right.removesuffix("cm"))
        tw = text_area_twips(page_config.size, left_cm, right_cm)
        center_pos = tw // 2
        right_pos = tw
    else:
        # Fallback: A4 with 3.17cm margins
        center_pos = 4153
        right_pos = 8306

    chapter_num = 0
    formula_count_in_chapter = 0
    continuous_count = 0

    for para_idx, paragraph in enumerate(document.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading 1"):
            chapter_num += 1
            formula_count_in_chapter = 0

        if not is_formula_paragraph(paragraph):
            continue

        if formula_config.numbering == "chapter":
            formula_count_in_chapter += 1
            label = f"({chapter_num}-{formula_count_in_chapter})"
        else:  # continuous
            continuous_count += 1
            label = f"({continuous_count})"

        # Set tab stops: center tab at midpoint, right tab at right margin
        _set_tab_stops(paragraph, center_pos, right_pos)
        # Paragraph alignment: left (tabs handle centering)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Insert leading tab BEFORE the formula element
        p = paragraph._p
        omath = p.find(qn("m:oMath"))
        if omath is not None:
            leading_tab = _make_run_with_tab()
            p.insert(list(p).index(omath), leading_tab)

        # Append trailing tab + number label with font
        p.append(_make_run_with_tab())
        label_run = OxmlElement("w:r")
        # Set font on the label run
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), formula_config.font_en)
        rFonts.set(qn("w:hAnsi"), formula_config.font_en)
        rFonts.set(qn("w:eastAsia"), formula_config.font_cn)
        rPr.append(rFonts)
        label_run.append(rPr)
        label_t = OxmlElement("w:t")
        label_t.text = label
        label_run.append(label_t)
        p.append(label_run)

        # Register with registry and add bookmark if provided
        if registry is not None and formula_items is not None:
            for fi_idx, fi_item in formula_items:
                if fi_idx == para_idx:
                    entry = registry.register("formula", fi_item["key"], label, para_idx)
                    from .references import wrap_in_bookmark
                    wrap_in_bookmark(p, label_run, entry)
                    break
