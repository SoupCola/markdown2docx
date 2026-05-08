from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from .styles import apply_font
from .utils import parse_distance

PAGE_NUMBER_ALIGNMENT = {
    "bottom_center": WD_ALIGN_PARAGRAPH.CENTER,
    "top_center": WD_ALIGN_PARAGRAPH.CENTER,
    "top_right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _apply_measure(value: str):
    unit, number = parse_distance(value)
    if unit == "cm":
        return Cm(number)
    if unit == "in":
        return Inches(number)
    return Pt(number)


def _add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def apply_page_settings(document, page_config):
    section = document.sections[0]
    section.top_margin = _apply_measure(page_config.margins.top)
    section.bottom_margin = _apply_measure(page_config.margins.bottom)
    section.left_margin = _apply_measure(page_config.margins.left)
    section.right_margin = _apply_measure(page_config.margins.right)

    header = section.header.paragraphs[0]
    header.text = page_config.header.content
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header.runs:
        apply_font(
            header.runs[0],
            page_config.header.font_cn,
            page_config.header.font_en,
            page_config.header.size,
            False,
            False,
        )

    footer = section.footer.paragraphs[0]
    footer.text = page_config.footer.content
    footer.alignment = PAGE_NUMBER_ALIGNMENT[page_config.page_number.position]
    _add_page_number(footer)
    # Apply font to all footer runs (text + page number)
    for run in footer.runs:
        apply_font(
            run,
            page_config.header.font_cn,
            page_config.header.font_en,
            page_config.header.size,
            False,
            False,
        )
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), str(page_config.page_number.start))
    section._sectPr.append(pg_num_type)
