from __future__ import annotations

from pathlib import Path

from docx import Document


def analyze(filepath: Path) -> dict:
    document = Document(filepath)
    heading_counts = {"level_1": 0, "level_2": 0, "level_3": 0}
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name == "Heading 1":
            heading_counts["level_1"] += 1
        if style_name == "Heading 2":
            heading_counts["level_2"] += 1
        if style_name == "Heading 3":
            heading_counts["level_3"] += 1
        paragraphs.append(
            {
                "index": index,
                "text": paragraph.text,
                "style": style_name,
                "font": paragraph.runs[0].font.name if paragraph.runs else None,
            }
        )

    sections = []
    for section in document.sections:
        sections.append(
            {
                "page_size": f"{section.page_width}x{section.page_height}",
                "margins": {
                    "top": section.top_margin.cm,
                    "bottom": section.bottom_margin.cm,
                    "left": section.left_margin.cm,
                    "right": section.right_margin.cm,
                },
            }
        )

    return {
        "paragraphs": paragraphs,
        "tables": [
            {"index": index, "rows": len(table.rows), "cols": len(table.columns), "has_caption": False}
            for index, table in enumerate(document.tables)
        ],
        "images": [],
        "sections": sections,
        "total_paragraphs": len(document.paragraphs),
        "headings_count": heading_counts,
    }


def summarize_analysis(result: dict) -> str:
    return (
        f"段落数: {result['total_paragraphs']}\n"
        f"一级标题: {result['headings_count']['level_1']}\n"
        f"二级标题: {result['headings_count']['level_2']}\n"
        f"三级标题: {result['headings_count']['level_3']}\n"
        f"表格数: {len(result['tables'])}"
    )
