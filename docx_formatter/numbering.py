from __future__ import annotations

from dataclasses import dataclass
import re
from typing import Any

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

LEVEL_1_PATTERN = re.compile(r"^(\d+)\.\s+(.+)$")
LEVEL_2_PATTERN = re.compile(r"^（(\d+)）(.+)$")
BRACKET_PATTERN = re.compile(r"^\[(\d+)\]\s*(.+)$")
FIGURE_LIKE_PATTERN = re.compile(r"^[图表]\s*\d+\.")


@dataclass
class NumberingCandidate:
    paragraph_index: int
    kind: str
    number: int
    prefix: str
    text: str
    paragraph: Any


@dataclass
class RejectedNumberingItem:
    paragraph_index: int
    kind: str
    number: int
    prefix: str
    text: str
    reason: str
    paragraph: Any


@dataclass
class DetectedNumberingAnalysis:
    convertible_items: list[NumberingCandidate]
    rejected_items: list[RejectedNumberingItem]


def _parse_candidate(paragraph, index: int) -> NumberingCandidate | None:
    text = paragraph.text.strip()
    if not text or FIGURE_LIKE_PATTERN.match(text):
        return None

    for kind, pattern in (
        ("level_1", LEVEL_1_PATTERN),
        ("level_2", LEVEL_2_PATTERN),
        ("bracket", BRACKET_PATTERN),
    ):
        match = pattern.match(text)
        if match:
            number = int(match.group(1))
            body = match.group(2).strip()
            return NumberingCandidate(
                paragraph_index=index,
                kind=kind,
                number=number,
                prefix=match.group(0)[: len(match.group(0)) - len(match.group(2))],
                text=body,
                paragraph=paragraph,
            )
    return None


def _evaluate_flat_sequence(
    candidates: list[NumberingCandidate],
    convertible_items: list[NumberingCandidate],
    rejected_items: list[RejectedNumberingItem],
) -> None:
    if not candidates:
        return
    if candidates[0].number != 1:
        reason = "sequence_must_start_at_1"
    else:
        expected = 1
        reason = None
        for candidate in candidates:
            if candidate.number != expected:
                reason = "sequence_must_be_contiguous"
                break
            expected += 1

    if reason is None:
        convertible_items.extend(candidates)
        return

    for candidate in candidates:
        rejected_items.append(
            RejectedNumberingItem(
                paragraph_index=candidate.paragraph_index,
                kind=candidate.kind,
                number=candidate.number,
                prefix=candidate.prefix,
                text=candidate.text,
                reason=reason,
                paragraph=candidate.paragraph,
            )
        )


def _append_rejected_group(
    candidates: list[NumberingCandidate],
    rejected_items: list[RejectedNumberingItem],
    reason: str,
) -> None:
    for candidate in candidates:
        rejected_items.append(
            RejectedNumberingItem(
                paragraph_index=candidate.paragraph_index,
                kind=candidate.kind,
                number=candidate.number,
                prefix=candidate.prefix,
                text=candidate.text,
                reason=reason,
                paragraph=candidate.paragraph,
            )
        )


def _flush_level_2_group(
    current_group: list[NumberingCandidate],
    convertible_items: list[NumberingCandidate],
    rejected_items: list[RejectedNumberingItem],
) -> None:
    if not current_group:
        return
    if current_group[0].number != 1:
        _append_rejected_group(current_group, rejected_items, "missing_parent_level_1_context")
        return

    expected = 1
    for candidate in current_group:
        if candidate.number != expected:
            _append_rejected_group(current_group, rejected_items, "sequence_must_be_contiguous")
            return
        expected += 1

    convertible_items.extend(current_group)


def _evaluate_level_2_sequences(
    candidates: list[NumberingCandidate],
    convertible_level_1: list[NumberingCandidate],
    convertible_items: list[NumberingCandidate],
    rejected_items: list[RejectedNumberingItem],
) -> None:
    if not candidates:
        return

    parent_indexes = {item.paragraph_index for item in convertible_level_1}
    current_group: list[NumberingCandidate] = []
    seen_parent = False
    last_index = -1

    for candidate in candidates:
        has_parent_before = any(parent_index < candidate.paragraph_index for parent_index in parent_indexes)
        if has_parent_before:
            if current_group and candidate.paragraph_index != last_index + 1:
                _flush_level_2_group(current_group, convertible_items, rejected_items)
                current_group = []
            seen_parent = True

        if not seen_parent:
            _append_rejected_group([candidate], rejected_items, "missing_parent_level_1_context")
            continue

        current_group.append(candidate)
        last_index = candidate.paragraph_index

    _flush_level_2_group(current_group, convertible_items, rejected_items)


def _split_into_sub_groups(
    candidates: list[NumberingCandidate], all_candidate_indices: set[int],
) -> list[list[NumberingCandidate]]:
    """Split candidates into sub-groups where gaps contain only other candidates."""
    if not candidates:
        return []
    groups: list[list[NumberingCandidate]] = []
    current = [candidates[0]]
    for i in range(1, len(candidates)):
        prev_idx = candidates[i - 1].paragraph_index
        curr_idx = candidates[i].paragraph_index
        gap_ok = all(j in all_candidate_indices for j in range(prev_idx + 1, curr_idx))
        if gap_ok:
            current.append(candidates[i])
        else:
            groups.append(current)
            current = [candidates[i]]
    groups.append(current)
    return groups


def detect_numbering_groups(document, numbering_config) -> DetectedNumberingAnalysis:
    del numbering_config

    convertible_items: list[NumberingCandidate] = []
    rejected_items: list[RejectedNumberingItem] = []
    grouped_candidates: dict[str, list[NumberingCandidate]] = {
        "level_1": [],
        "level_2": [],
        "bracket": [],
    }

    for index, paragraph in enumerate(document.paragraphs):
        candidate = _parse_candidate(paragraph, index)
        if candidate is None:
            continue
        grouped_candidates[candidate.kind].append(candidate)

    all_candidate_indices = {
        c.paragraph_index
        for group in grouped_candidates.values()
        for c in group
    }

    for sub in _split_into_sub_groups(grouped_candidates["level_1"], all_candidate_indices):
        _evaluate_flat_sequence(sub, convertible_items, rejected_items)
    for sub in _split_into_sub_groups(grouped_candidates["bracket"], all_candidate_indices):
        _evaluate_flat_sequence(sub, convertible_items, rejected_items)

    convertible_level_1 = [item for item in convertible_items if item.kind == "level_1"]
    _evaluate_level_2_sequences(
        grouped_candidates["level_2"],
        convertible_level_1,
        convertible_items,
        rejected_items,
    )

    convertible_items.sort(key=lambda item: item.paragraph_index)
    return DetectedNumberingAnalysis(
        convertible_items=convertible_items,
        rejected_items=rejected_items,
    )


def _make_level(ilvl: int, num_fmt: str, lvl_text: str):
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), str(ilvl))
    lvl.set(qn("w:tentative"), "1")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), lvl_text)
    lvl.append(text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl.append(suff)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    return lvl


def _make_num(num_id: int, abstract_id: int):
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    return num


def _register_numbering_definitions(document) -> dict[str, tuple[int, int]]:
    numbering = document.part.numbering_part.element
    existing_abstract = numbering.findall(qn("w:abstractNum"))
    next_abstract_id = max(
        (int(a.get(qn("w:abstractNumId"))) for a in existing_abstract), default=-1
    ) + 1
    existing_nums = numbering.findall(qn("w:num"))
    next_num_id = max(
        (int(n.get(qn("w:numId"))) for n in existing_nums), default=0
    ) + 1

    abstract_multi = OxmlElement("w:abstractNum")
    abstract_multi.set(qn("w:abstractNumId"), str(next_abstract_id))
    abstract_multi.append(_make_level(0, "decimal", "%1."))
    abstract_multi.append(_make_level(1, "decimal", "（%2）"))
    numbering.insert(0, abstract_multi)

    abstract_bracket = OxmlElement("w:abstractNum")
    abstract_bracket.set(qn("w:abstractNumId"), str(next_abstract_id + 1))
    abstract_bracket.append(_make_level(0, "decimal", "[%1]"))
    numbering.insert(1, abstract_bracket)

    numbering.append(_make_num(next_num_id, next_abstract_id))
    numbering.append(_make_num(next_num_id + 1, next_abstract_id + 1))

    return {
        "numbered_level_1": (next_num_id, 0),
        "numbered_level_2": (next_num_id, 1),
        "numbered_bracket": (next_num_id + 1, 0),
    }


def _apply_numPr(paragraph, num_id: int, ilvl: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numPr.append(ilvl_el)
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(numId_el)
    pPr.insert(0, numPr)


def apply_generated_numbering(document, paragraph_specs, numbering_config) -> None:
    numbered_items = [
        (i, s) for i, s in enumerate(paragraph_specs) if s["type"].startswith("numbered_")
    ]
    if not numbered_items:
        return
    for _, spec in numbered_items:
        document.add_paragraph(spec["text"])
    num_map = _register_numbering_definitions(document)
    start_idx = len(document.paragraphs) - len(numbered_items)
    for offset, (_, spec) in enumerate(numbered_items):
        num_id, ilvl = num_map[spec["type"]]
        _apply_numPr(document.paragraphs[start_idx + offset], num_id, ilvl)


_KIND_TO_NUM_KEY = {
    "level_1": "numbered_level_1",
    "level_2": "numbered_level_2",
    "bracket": "numbered_bracket",
}


def _replace_paragraph_text(paragraph, text: str) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.add_run(text)


def convert_detected_numbering(document, analysis: DetectedNumberingAnalysis, numbering_config) -> None:
    if not analysis.convertible_items:
        return
    num_map = _register_numbering_definitions(document)
    for item in analysis.convertible_items:
        num_key = _KIND_TO_NUM_KEY[item.kind]
        num_id, ilvl = num_map[num_key]
        _replace_paragraph_text(item.paragraph, item.text)
        _apply_numPr(item.paragraph, num_id, ilvl)
