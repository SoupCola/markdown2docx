from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .config import FormatterConfig, load_template, merge_overrides
from .template_paths import default_template_path

_DEFAULT_TEMPLATE = default_template_path()
_CN_SIZES = {
    42.0: "初号",
    36.0: "小初",
    26.0: "一号",
    24.0: "小一",
    22.0: "二号",
    18.0: "小二",
    16.0: "三号",
    15.0: "小三",
    14.0: "四号",
    12.0: "小四",
    10.5: "五号",
    9.0: "小五",
}
_ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
    None: "left",
}


def _base_config() -> FormatterConfig:
    return load_template(_DEFAULT_TEMPLATE)


def _pt_to_size_name(value_pt: float | None, fallback: str) -> str:
    if value_pt is None:
        return fallback
    rounded = round(value_pt, 1)
    return _CN_SIZES.get(rounded, f"{rounded:g}pt")


def _font_cn(style, fallback: str) -> str:
    rpr = style._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        east = rpr.rFonts.get(qn("w:eastAsia"))
        if east:
            return east
    return fallback


def _font_en(style, fallback: str) -> str:
    return style.font.name or fallback


def _cm_string(cm_value: float | None, fallback: str) -> str:
    if cm_value is None:
        return fallback
    return f"{cm_value:.2f}cm"


def _alignment_name(value, fallback: str) -> str:
    return _ALIGNMENT_MAP.get(value, fallback)


def _indent_string(first_line_indent, fallback: str) -> str:
    if first_line_indent is None:
        return fallback
    pt = first_line_indent.pt
    if abs(pt - 24) < 0.2:
        return "2chars"
    return f"{pt:g}pt"


def _line_spacing_string(line_spacing, fallback: str) -> str:
    if line_spacing is None:
        return fallback
    if isinstance(line_spacing, (int, float)):
        return f"{line_spacing:g}lines"
    return fallback


def _extract_style_overrides(document) -> dict:
    base = _base_config()
    heading_1 = document.styles["Heading 1"]
    heading_2 = document.styles["Heading 2"]
    heading_3 = document.styles["Heading 3"]
    normal = document.styles["Normal"]
    section = document.sections[0]
    header_text = ""
    if section.header.paragraphs:
        header_text = " ".join(p.text for p in section.header.paragraphs).strip()

    return {
        "page": {
            "size": "A4" if abs(section.page_width.cm - 21.0) < 0.3 else base.page.size,
            "margins": {
                "top": _cm_string(section.top_margin.cm, base.page.margins.top),
                "bottom": _cm_string(section.bottom_margin.cm, base.page.margins.bottom),
                "left": _cm_string(section.left_margin.cm, base.page.margins.left),
                "right": _cm_string(section.right_margin.cm, base.page.margins.right),
            },
            "header": {
                "content": header_text or base.page.header.content,
            },
        },
        "headings": {
            "level_1": {
                "font_cn": _font_cn(heading_1, base.headings.level_1.font_cn),
                "font_en": _font_en(heading_1, base.headings.level_1.font_en),
                "size": _pt_to_size_name(heading_1.font.size.pt if heading_1.font.size else None, base.headings.level_1.size),
                "bold": bool(heading_1.font.bold) if heading_1.font.bold is not None else base.headings.level_1.bold,
                "align": _alignment_name(heading_1.paragraph_format.alignment, base.headings.level_1.align),
            },
            "level_2": {
                "font_cn": _font_cn(heading_2, base.headings.level_2.font_cn),
                "font_en": _font_en(heading_2, base.headings.level_2.font_en),
                "size": _pt_to_size_name(heading_2.font.size.pt if heading_2.font.size else None, base.headings.level_2.size),
                "bold": bool(heading_2.font.bold) if heading_2.font.bold is not None else base.headings.level_2.bold,
                "align": _alignment_name(heading_2.paragraph_format.alignment, base.headings.level_2.align),
            },
            "level_3": {
                "font_cn": _font_cn(heading_3, base.headings.level_3.font_cn),
                "font_en": _font_en(heading_3, base.headings.level_3.font_en),
                "size": _pt_to_size_name(heading_3.font.size.pt if heading_3.font.size else None, base.headings.level_3.size),
                "bold": bool(heading_3.font.bold) if heading_3.font.bold is not None else base.headings.level_3.bold,
                "align": _alignment_name(heading_3.paragraph_format.alignment, base.headings.level_3.align),
            },
        },
        "body": {
            "font_cn": _font_cn(normal, base.body.font_cn),
            "font_en": _font_en(normal, base.body.font_en),
            "size": _pt_to_size_name(normal.font.size.pt if normal.font.size else None, base.body.size),
            "first_line_indent": _indent_string(normal.paragraph_format.first_line_indent, base.body.first_line_indent),
            "line_spacing": _line_spacing_string(normal.paragraph_format.line_spacing, base.body.line_spacing),
            "align": _alignment_name(normal.paragraph_format.alignment, base.body.align),
        },
    }


def _styles_are_informative(document) -> bool:
    normal = document.styles["Normal"]
    heading_1 = document.styles["Heading 1"]
    text_spec_overrides = _parse_text_spec(document)
    if text_spec_overrides:
        section = document.sections[0]
        default_margins = _base_config().page.margins
        margin_overridden = any(
            (
                section.top_margin is not None and abs(section.top_margin.cm - float(default_margins.top.removesuffix("cm"))) > 0.01,
                section.bottom_margin is not None and abs(section.bottom_margin.cm - float(default_margins.bottom.removesuffix("cm"))) > 0.01,
                section.left_margin is not None and abs(section.left_margin.cm - float(default_margins.left.removesuffix("cm"))) > 0.01,
                section.right_margin is not None and abs(section.right_margin.cm - float(default_margins.right.removesuffix("cm"))) > 0.01,
            )
        )
        return margin_overridden
    return bool(normal.font.name or normal.font.size or heading_1.font.name or heading_1.font.size)


def _parse_text_spec(document) -> dict:
    overrides: dict = {}
    text = "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())

    heading_match = re.search(r"标题(?:（一）)?用(?P<font>黑体|宋体|楷体)(?P<size>[初小一二三四五六七八号]+|小[一二三四五六])(?P<align>居中|左对齐|右对齐|两端对齐)?", text)
    if heading_match:
        align_map = {"居中": "center", "左对齐": "left", "右对齐": "right", "两端对齐": "justify", None: "center"}
        overrides.setdefault("headings", {})["level_1"] = {
            "font_cn": heading_match.group("font"),
            "size": heading_match.group("size"),
            "align": align_map[heading_match.group("align")],
        }

    body_match = re.search(r"正文用(?P<font>黑体|宋体|楷体)(?P<size>[初小一二三四五六七八号]+|小[一二三四五六])", text)
    if body_match:
        overrides.setdefault("body", {}).update({
            "font_cn": body_match.group("font"),
            "size": body_match.group("size"),
        })

    indent_match = re.search(r"首行缩进(?P<count>\d+)字符", text)
    if indent_match:
        overrides.setdefault("body", {})["first_line_indent"] = f"{indent_match.group('count')}chars"

    spacing_match = re.search(r"行距(?P<value>\d+(?:\.\d+)?)倍", text)
    if spacing_match:
        overrides.setdefault("body", {})["line_spacing"] = f"{spacing_match.group('value')}lines"

    margin_match = re.search(r"上下(?P<tb>\d+(?:\.\d+)?)cm[，,]?左右(?P<lr>\d+(?:\.\d+)?)cm", text)
    if margin_match:
        overrides.setdefault("page", {}).setdefault("margins", {}).update({
            "top": f"{margin_match.group('tb')}cm",
            "bottom": f"{margin_match.group('tb')}cm",
            "left": f"{margin_match.group('lr')}cm",
            "right": f"{margin_match.group('lr')}cm",
        })

    header_match = re.search(r"页眉内容[:：](?P<content>[^\n]+)", text)
    if header_match:
        overrides.setdefault("page", {}).setdefault("header", {})["content"] = header_match.group("content").strip()

    return overrides


def extract_format_from_docx(docx_path: Path) -> FormatterConfig:
    document = Document(docx_path)
    base = _base_config()
    overrides = _extract_style_overrides(document) if _styles_are_informative(document) else _parse_text_spec(document)
    if _styles_are_informative(document):
        return merge_overrides(base, overrides)
    return merge_overrides(base, _parse_text_spec(document))
