from __future__ import annotations

import re
from dataclasses import dataclass, field

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styles import ALIGNMENT_MAP, apply_font
from .utils import cn_size_to_pt, parse_indent


@dataclass
class LabelEntry:
    kind: str           # "figure" | "table" | "formula" | "reference"
    key: str            # user-specified unique ID, e.g. "fig:system_arch"
    label_text: str     # display text, e.g. "图 1-1" or "(1-1)" or "[1]"
    paragraph_index: int
    bookmark_name: str  # e.g. "_fig_system_arch"
    bookmark_id: int    # auto-incrementing


@dataclass
class LabelRegistry:
    entries: dict[str, LabelEntry] = field(default_factory=dict)
    _next_bookmark_id: int = 1

    def register(self, kind: str, key: str, label_text: str, paragraph_index: int) -> LabelEntry:
        bookmark_name = f"_{key.replace(':', '_')}"
        entry = LabelEntry(
            kind=kind,
            key=key,
            label_text=label_text,
            paragraph_index=paragraph_index,
            bookmark_name=bookmark_name,
            bookmark_id=self._next_bookmark_id,
        )
        self.entries[key] = entry
        self._next_bookmark_id += 1
        return entry

    def lookup(self, key: str) -> LabelEntry | None:
        return self.entries.get(key)


def wrap_in_bookmark(paragraph_element, run_element, entry: LabelEntry) -> None:
    """Wrap a run element in w:bookmarkStart and w:bookmarkEnd."""
    p = paragraph_element
    r = run_element
    idx = list(p).index(r)

    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(entry.bookmark_id))
    bm_start.set(qn("w:name"), entry.bookmark_name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(entry.bookmark_id))

    p.insert(idx, bm_start)
    p.insert(idx + 2, bm_end)


def make_ref_field(bookmark_name: str, display_text: str, superscript: bool = False,
                   font_en: str | None = None, font_cn: str | None = None) -> list:
    """Create a list of XML elements for a Word REF field code.

    Returns 5 elements: begin, instrText, separate, display text run, end.
    If superscript is True, the display text run uses superscript vertical alignment.
    """
    begin_r = OxmlElement("w:r")
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin_r.append(begin_char)

    instr_r = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" REF {bookmark_name} \\h "
    instr_r.append(instr_text)

    sep_r = OxmlElement("w:r")
    sep_char = OxmlElement("w:fldChar")
    sep_char.set(qn("w:fldCharType"), "separate")
    sep_r.append(sep_char)

    display_r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if font_en:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font_en)
        rFonts.set(qn("w:hAnsi"), font_en)
        if font_cn:
            rFonts.set(qn("w:eastAsia"), font_cn)
        rPr.append(rFonts)
    if superscript:
        vertAlign = OxmlElement("w:vertAlign")
        vertAlign.set(qn("w:val"), "superscript")
        rPr.append(vertAlign)
    if len(rPr) > 0:
        display_r.append(rPr)
    display_t = OxmlElement("w:t")
    display_t.text = display_text
    display_r.append(display_t)

    end_r = OxmlElement("w:r")
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end_r.append(end_char)

    return [begin_r, instr_r, sep_r, display_r, end_r]


_PLACEHOLDER_RE = re.compile(r"\{([A-Za-z]+:[^}]+)\}")


def parse_placeholders(text: str) -> list[tuple[str, str]]:
    """Parse text with {key} placeholders into segments.

    Returns a list of ("text", "plain text") and ("ref", "key") tuples.
    """
    segments = []
    last_end = 0
    for match in _PLACEHOLDER_RE.finditer(text):
        start, end = match.span()
        if start > last_end:
            segments.append(("text", text[last_end:start]))
        segments.append(("ref", match.group(1)))
        last_end = end
    if last_end < len(text):
        segments.append(("text", text[last_end:]))
    return segments


def _number_label(index: int, style: str) -> str:
    if style == "bracket":
        return f"[{index}]"
    if style == "paren":
        return f"({index})"
    return f"{index}."


def _make_seq_field(seq_name: str, display_text: str, font_en: str | None = None, font_cn: str | None = None) -> list:
    """Create a SEQ field that auto-increments. Returns XML element list.

    Example output: { SEQ ref \\* ARABIC } → display_text
    """
    begin_r = OxmlElement("w:r")
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin_r.append(begin_char)

    instr_r = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" SEQ {seq_name} \\* ARABIC "
    instr_r.append(instr_text)

    sep_r = OxmlElement("w:r")
    sep_char = OxmlElement("w:fldChar")
    sep_char.set(qn("w:fldCharType"), "separate")
    sep_r.append(sep_char)

    display_r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if font_en:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font_en)
        rFonts.set(qn("w:hAnsi"), font_en)
        if font_cn:
            rFonts.set(qn("w:eastAsia"), font_cn)
        rPr.append(rFonts)
    display_r.append(rPr)
    display_t = OxmlElement("w:t")
    display_t.text = display_text
    display_r.append(display_t)

    end_r = OxmlElement("w:r")
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end_r.append(end_char)

    return [begin_r, instr_r, sep_r, display_r, end_r]


def _wrap_elements_in_bookmark(paragraph_element, elements: list, entry: LabelEntry) -> None:
    """Wrap multiple XML elements in a bookmark pair."""
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(entry.bookmark_id))
    bm_start.set(qn("w:name"), entry.bookmark_name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(entry.bookmark_id))

    # Insert bookmarkStart before the first element
    first = elements[0]
    idx = list(paragraph_element).index(first)
    paragraph_element.insert(idx, bm_start)

    # Insert bookmarkEnd after the last element
    last = elements[-1]
    idx = list(paragraph_element).index(last)
    paragraph_element.insert(idx + 1, bm_end)


def generate_bibliography(document, registry: LabelRegistry, references: list[dict], config) -> None:
    """Generate bibliography section at end of document with bookmarks and SEQ auto-numbering."""
    title_p = document.add_paragraph(config.title)
    title_p.alignment = ALIGNMENT_MAP.get(config.title_align, WD_ALIGN_PARAGRAPH.LEFT)
    if title_p.runs:
        apply_font(
            title_p.runs[0],
            config.title_font_cn, config.title_font_en,
            config.title_size, config.title_bold, False,
        )

    for i, ref in enumerate(references):
        index = i + 1
        entry = registry.lookup(ref["key"])
        if entry is None:
            continue

        label_text = _number_label(index, config.number_style)
        entry.label_text = label_text

        p = document.add_paragraph()
        p.alignment = ALIGNMENT_MAP.get(config.align, WD_ALIGN_PARAGRAPH.LEFT)

        # Build bracket + SEQ field for auto-numbering
        # Format: [ {SEQ ref} ] space  reference_text
        bracket_open = p.add_run("[")
        apply_font(bracket_open, config.font_cn, config.font_en, config.size, False, False)

        # SEQ field for the number
        seq_elements = _make_seq_field(
            "ref", str(index),
            font_en=config.font_en, font_cn=config.font_cn,
        )
        for el in seq_elements:
            p._p.append(el)

        bracket_close = p.add_run("] ")
        apply_font(bracket_close, config.font_cn, config.font_en, config.size, False, False)

        entry.paragraph_index = len(document.paragraphs) - 1
        # Wrap bracket_open run + SEQ elements + bracket_close run in bookmark
        bookmark_elements = [bracket_open._r] + seq_elements + [bracket_close._r]
        _wrap_elements_in_bookmark(p._p, bookmark_elements, entry)

        text_run = p.add_run(ref["text"])
        apply_font(text_run, config.font_cn, config.font_en, config.size, False, False)

        if config.hanging_indent:
            font_size = cn_size_to_pt(config.size)
            indent_pt = parse_indent(config.hanging_indent, font_size)
            p.paragraph_format.first_line_indent = Pt(-indent_pt)
            p.paragraph_format.left_indent = Pt(indent_pt)
