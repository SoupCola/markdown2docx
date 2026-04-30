from __future__ import annotations

import re

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .references import wrap_in_bookmark
from .styles import apply_font

FIGURE_PATTERN = re.compile(r"^图\s+\d+(?:-\d+)?")
TABLE_PATTERN = re.compile(r"^表\s+\d+(?:-\d+)?")


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, edge_data in kwargs.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def apply_figure_table_formatting(document, figures_config):
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if FIGURE_PATTERN.match(text) or TABLE_PATTERN.match(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                apply_font(
                    run,
                    figures_config.caption_font_cn,
                    figures_config.caption_font_en,
                    figures_config.caption_size,
                    figures_config.caption_bold,
                    False,
                )
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        if row_index == 0:
                            run.bold = True
        if figures_config.table_style == "three_line":
            for cell in table.rows[0].cells:
                _set_cell_border(
                    cell,
                    top={"val": "single", "sz": 12},
                    bottom={"val": "single", "sz": 6},
                )
            for cell in table.rows[-1].cells:
                _set_cell_border(cell, bottom={"val": "single", "sz": 12})


def auto_number_captions(document, registry, kind: str, items: list[tuple[int, dict]], config) -> None:
    """Auto-number figure or table captions and register them with bookmarks.

    Args:
        document: python-docx Document object
        registry: LabelRegistry to register labels into
        kind: "figure" or "table"
        items: list of (paragraph_index, item_dict) where item_dict has "key" and "caption"
        config: FiguresConfig with numbering mode
    """
    prefix = "图" if kind == "figure" else "表"

    for item_idx, (paragraph_index, item) in enumerate(items):
        paragraph = document.paragraphs[paragraph_index]
        caption_text = item.get("caption", "")

        if config.numbering == "chapter":
            # Count Heading 1 paragraphs before this index
            chapter_num = 0
            for i in range(paragraph_index):
                p = document.paragraphs[i]
                style_name = p.style.name if p.style else ""
                if style_name.startswith("Heading 1"):
                    chapter_num += 1
            # Count items of same kind in same chapter that come before this one
            count_in_chapter = 0
            for pi, _ in items[:item_idx]:
                if pi >= paragraph_index:
                    break
                ch = 0
                for i in range(pi):
                    p = document.paragraphs[i]
                    style_name = p.style.name if p.style else ""
                    if style_name.startswith("Heading 1"):
                        ch += 1
                if ch == chapter_num:
                    count_in_chapter += 1
            count_in_chapter += 1
            label_text = f"{prefix} {chapter_num}-{count_in_chapter}"
        else:
            label_text = f"{prefix} {item_idx + 1}"

        paragraph.clear()
        if caption_text:
            run = paragraph.add_run(f"{label_text} {caption_text}")
        else:
            run = paragraph.add_run(label_text)

        entry = registry.register(kind, item["key"], label_text, paragraph_index)
        wrap_in_bookmark(paragraph._p, run._r, entry)
