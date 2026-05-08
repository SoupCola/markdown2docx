from __future__ import annotations

from pathlib import Path
from dataclasses import asdict

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .analyzer import analyze, summarize_analysis
from .config import load_template, merge_overrides
from .figures import apply_figure_table_formatting, auto_number_captions
from .format_extractor import extract_format_from_docx
from .formulas import apply_formula_formatting, apply_formula_numbering, insert_formula_paragraph
from .headings import apply_heading_formatting
from .md_parser import parse_markdown_files
from .numbering import _apply_numPr, _register_numbering_definitions, convert_detected_numbering, detect_numbering_groups
from .pages import apply_page_settings
from .paragraphs import apply_body_formatting
from .references import LabelRegistry, parse_placeholders, make_ref_field, generate_bibliography
from .styles import apply_font


REFERENCE_HEADING_TITLES = {"参考文献"}


def _resolved_config(template_path: Path, overrides: dict):
    config = load_template(template_path)
    return merge_overrides(config, overrides)


def _insert_image(document, item: dict, page_config=None) -> None:
    """Insert an image from local path, centered, with optional caption.

    Auto-scales to fit page text area if image is too wide.
    """
    from PIL import Image
    from .utils import parse_distance

    image_path = Path(item["path"])
    width = item.get("width")

    if width is None:
        # Auto-scale: get image natural width in inches
        img = Image.open(str(image_path))
        dpi = img.info.get("dpi", (96, 96))
        natural_width_in = img.width / dpi[0]
        img.close()

        # Calculate max width from page text area
        if page_config is not None:
            page_widths_mm = {"A4": 210, "Letter": 216}
            pw_mm = page_widths_mm.get(page_config.size, 210)
            left_u, left_v = parse_distance(page_config.margins.left)
            right_u, right_v = parse_distance(page_config.margins.right)
            left_mm = left_v * 10 if left_u == "cm" else left_v * (10 if left_u == "mm" else 25.4)
            right_mm = right_v * 10 if right_u == "cm" else right_v * (10 if right_u == "mm" else 25.4)
            text_width_in = (pw_mm - left_mm - right_mm) / 25.4
        else:
            text_width_in = 5.77  # A4 with 3.17cm margins

        # Use natural width if smaller than text area, otherwise cap at text area
        width = min(natural_width_in, text_width_in)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    # Add caption paragraph if provided
    caption = item.get("caption")
    if caption:
        cap_paragraph = document.add_paragraph(caption)
        cap_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _insert_table(document, item: dict) -> int | None:
    rows = item.get("rows", [])
    if not rows:
        return None

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row in enumerate(rows):
        for col_idx in range(column_count):
            value = row[col_idx] if col_idx < len(row) else ""
            table.cell(row_idx, col_idx).text = value

    caption = item.get("caption")
    if caption is None:
        return None

    cap_paragraph = document.add_paragraph(caption)
    cap_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return len(document.paragraphs) - 1


def format_document(source_path: Path, template_path: Path, overrides: dict) -> Path:
    document = Document(source_path)
    config = _resolved_config(template_path, overrides)
    apply_page_settings(document, config.page)
    apply_heading_formatting(document, config.headings)
    apply_body_formatting(document, config.body)
    apply_figure_table_formatting(document, config.figures)
    analysis = detect_numbering_groups(document, config.numbering)
    convert_detected_numbering(document, analysis, config.numbering)
    for item in analysis.convertible_items:
        item.paragraph.paragraph_format.first_line_indent = Pt(0)
    apply_formula_formatting(document, config.formulas)
    apply_formula_numbering(document, config.formulas, config.page)
    output_path = source_path.with_name(f"{source_path.stem}_formatted.docx")
    document.save(output_path)
    return output_path


def create_from_markdown(
    output_path: Path,
    md_paths: list[Path],
    format_docx_path: Path | None = None,
    template_path: Path | None = None,
    overrides: dict | None = None,
) -> Path:
    paragraphs, references = parse_markdown_files(md_paths)
    resolved_overrides = dict(overrides or {})
    resolved_template_path = template_path

    if format_docx_path is not None:
        extracted = extract_format_from_docx(format_docx_path)
        if resolved_template_path is None:
            raise ValueError("template_path is required when format_docx_path is provided")
        resolved_config = merge_overrides(load_template(resolved_template_path), asdict(extracted))
        resolved_config = merge_overrides(resolved_config, resolved_overrides)
        extracted_overrides = asdict(resolved_config)
        extracted_overrides["headings"]["level_2"]["size"] = resolved_config.headings.level_2.size
        extracted_overrides["headings"]["level_3"]["size"] = resolved_config.headings.level_3.size
        return create_document(
            output_path=output_path,
            template_path=resolved_template_path,
            paragraphs=paragraphs,
            overrides=extracted_overrides,
            references=references,
        )

    if resolved_template_path is None:
        raise ValueError("template_path is required")

    return create_document(
        output_path=output_path,
        template_path=resolved_template_path,
        paragraphs=paragraphs,
        overrides=resolved_overrides,
        references=references,
    )


def create_document(output_path: Path, template_path: Path, paragraphs: list[dict], overrides: dict, references: list[dict] | None = None) -> Path:
    document = Document()
    config = _resolved_config(template_path, overrides)

    # Set default font on Normal style as fallback
    normal_style = document.styles['Normal']
    normal_style.font.name = config.body.font_en
    normal_rpr = normal_style._element.get_or_add_rPr()
    normal_rfonts = normal_rpr.find(qn('w:rFonts'))
    if normal_rfonts is None:
        normal_rfonts = OxmlElement('w:rFonts')
        normal_rpr.insert(0, normal_rfonts)
    normal_rfonts.set(qn('w:ascii'), config.body.font_en)
    normal_rfonts.set(qn('w:hAnsi'), config.body.font_en)
    normal_rfonts.set(qn('w:eastAsia'), config.body.font_cn)

    registry = LabelRegistry()
    references = list(references or [])
    citation_order: list[str] = []

    # Pre-register bibliography references so citations can resolve
    if references:
        for i, ref in enumerate(references):
            label = f"[{i + 1}]"
            registry.register("reference", ref["key"], label, -1)

    numbered_groups: list[list[tuple[int, dict]]] = []
    current_group: list[tuple[int, dict]] = []

    def _flush_group():
        nonlocal current_group
        if current_group:
            numbered_groups.append(current_group)
            current_group = []

    formula_indices: set[int] = set()
    formula_items: list[tuple[int, dict]] = []
    figure_items: list[tuple[int, dict]] = []
    table_items: list[tuple[int, dict]] = []
    # Track paragraphs with {key} placeholders for deferred resolution
    xref_paragraphs: list[tuple[int, str]] = []  # (paragraph_index, original_text)

    for item in paragraphs:
        if item["type"] == "heading_1":
            if item["text"].strip() in REFERENCE_HEADING_TITLES and references:
                _flush_group()
                continue
            document.add_heading(item["text"], level=1)
            _flush_group()
        elif item["type"] == "heading_2":
            document.add_heading(item["text"], level=2)
            _flush_group()
        elif item["type"] == "heading_3":
            document.add_heading(item["text"], level=3)
            _flush_group()
        elif item["type"].startswith("numbered_"):
            document.add_paragraph(item["text"])
            current_group.append((len(document.paragraphs) - 1, item))
        elif item["type"] == "formula":
            insert_formula_paragraph(document, item["latex"], config.formulas.font_en)
            para_idx = len(document.paragraphs) - 1
            formula_indices.add(para_idx)
            if "key" in item:
                formula_items.append((para_idx, item))
            _flush_group()
        elif item["type"] == "image":
            _insert_image(document, item, config.page)
            if "key" in item and item.get("caption"):
                caption_idx = len(document.paragraphs) - 1
                figure_items.append((caption_idx, item))
            _flush_group()
        elif item["type"] == "table":
            caption_idx = _insert_table(document, item)
            if "key" in item and caption_idx is not None:
                table_items.append((caption_idx, item))
            _flush_group()
        else:
            # Body text — store as plain text for now, resolve later
            text = item.get("text", "")
            segments = parse_placeholders(text)
            for seg_type, seg_text in segments:
                if seg_type == "ref" and seg_text.startswith("ref:") and seg_text not in citation_order:
                    citation_order.append(seg_text)
            has_refs = any(seg[0] == "ref" for seg in segments)
            if has_refs:
                p = document.add_paragraph()
                p.add_run(text)  # store full text; will be replaced in _resolve_xrefs
                xref_paragraphs.append((len(document.paragraphs) - 1, text))
            else:
                document.add_paragraph(text)
            _flush_group()
    _flush_group()

    numbered_indices: set[int] = set()
    for group in numbered_groups:
        num_map = _register_numbering_definitions(document)
        for doc_idx, spec in group:
            num_id, ilvl = num_map[spec["type"]]
            _apply_numPr(document.paragraphs[doc_idx], num_id, ilvl)
            numbered_indices.add(doc_idx)

    # Auto-number figure/table captions
    if figure_items:
        auto_number_captions(document, registry, "figure", figure_items, config.figures)
    if table_items:
        auto_number_captions(document, registry, "table", table_items, config.figures)

    apply_page_settings(document, config.page)
    apply_heading_formatting(document, config.headings)
    apply_body_formatting(document, config.body)
    apply_figure_table_formatting(document, config.figures)
    apply_formula_formatting(document, config.formulas)
    apply_formula_numbering(document, config.formulas, config.page, formula_items, registry)

    for idx in numbered_indices | formula_indices:
        document.paragraphs[idx].paragraph_format.first_line_indent = Pt(0)

    # Generate bibliography section
    bib_start_idx = len(document.paragraphs)
    if references:
        reference_by_key = {ref["key"]: ref for ref in references}
        ordered_references = [reference_by_key[key] for key in citation_order if key in reference_by_key]
        ordered_reference_keys = {ref["key"] for ref in ordered_references}
        ordered_references.extend(ref for ref in references if ref["key"] not in ordered_reference_keys)
        generate_bibliography(document, registry, ordered_references, config.references)
        # Clear indent on bibliography paragraphs
        for idx in range(bib_start_idx, len(document.paragraphs)):
            document.paragraphs[idx].paragraph_format.first_line_indent = Pt(0)

    # Resolve cross-references: replace {key} with REF fields
    for para_idx, original_text in xref_paragraphs:
        paragraph = document.paragraphs[para_idx]
        segments = parse_placeholders(original_text)
        # Clear existing runs
        for child in list(paragraph._p):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("r",):
                paragraph._p.remove(child)
        # Rebuild with text runs and REF fields
        for seg_type, seg_text in segments:
            if seg_type == "text":
                run = paragraph.add_run(seg_text)
                apply_font(run, config.body.font_cn, config.body.font_en, config.body.size, False, False)
            else:
                entry = registry.lookup(seg_text)
                if entry is not None and entry.label_text:
                    is_ref = seg_text.startswith("ref:")
                    field_elements = make_ref_field(
                        entry.bookmark_name, entry.label_text, superscript=is_ref,
                        font_en=config.body.font_en, font_cn=config.body.font_cn,
                    )
                    for el in field_elements:
                        paragraph._p.append(el)
                else:
                    run = paragraph.add_run("{" + seg_text + "}")
                    apply_font(run, config.body.font_cn, config.body.font_en, config.body.size, False, False)

    document.save(output_path)
    return output_path
