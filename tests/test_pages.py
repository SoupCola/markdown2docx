from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_formatter.config import load_template
from docx_formatter.pages import apply_page_settings


def test_apply_page_settings_updates_margins_header_footer_and_page_number():
    source_template = Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml")
    config = load_template(source_template)
    document = Document()

    apply_page_settings(document, config.page)

    section = document.sections[0]
    assert round(section.top_margin.cm, 2) == 2.54
    assert round(section.left_margin.cm, 2) == 3.17
    assert section.header.paragraphs[0].text == "毕业论文"
    footer_xml = section.footer._element.xml
    assert "PAGE" in footer_xml


def test_apply_page_settings_centers_header_content():
    source_template = Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml")
    config = load_template(source_template)
    document = Document()

    apply_page_settings(document, config.page)

    assert document.sections[0].header.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
