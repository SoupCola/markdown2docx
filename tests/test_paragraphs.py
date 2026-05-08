from dataclasses import replace
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_formatter.config import load_template
from docx_formatter.paragraphs import apply_body_formatting


def test_apply_body_formatting_updates_normal_paragraphs_only():
    document = Document()
    heading = document.add_heading("第一章 绪论", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("正文内容")
    blank = document.add_paragraph("   ")
    config = load_template(Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"))

    apply_body_formatting(document, config.body)

    assert heading.paragraph_format.first_line_indent is None
    assert round(paragraph.paragraph_format.first_line_indent.pt, 1) == 24.0
    assert paragraph.paragraph_format.alignment is not None
    assert run.font.name == "Times New Roman"
    assert blank.paragraph_format.first_line_indent is None


def test_apply_body_formatting_supports_distribute_alignment():
    document = Document()
    paragraph = document.add_paragraph("正文内容")
    config = load_template(Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"))
    distributed = replace(config.body, align="distribute")

    apply_body_formatting(document, distributed)

    assert paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.DISTRIBUTE
