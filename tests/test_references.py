from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def test_registry_register_returns_entry():
    from docx_formatter.references import LabelRegistry, LabelEntry

    registry = LabelRegistry()
    entry = registry.register("figure", "fig:arch", "图 1-1", 5)
    assert entry.kind == "figure"
    assert entry.key == "fig:arch"
    assert entry.label_text == "图 1-1"
    assert entry.paragraph_index == 5
    assert entry.bookmark_name == "_fig_arch"
    assert entry.bookmark_id == 1


def test_registry_register_auto_increments_id():
    from docx_formatter.references import LabelRegistry

    registry = LabelRegistry()
    e1 = registry.register("figure", "fig:a", "图 1-1", 0)
    e2 = registry.register("formula", "formula:e", "(1-1)", 1)
    assert e1.bookmark_id == 1
    assert e2.bookmark_id == 2


def test_registry_lookup_returns_entry():
    from docx_formatter.references import LabelRegistry

    registry = LabelRegistry()
    registry.register("figure", "fig:arch", "图 1-1", 5)
    entry = registry.lookup("fig:arch")
    assert entry is not None
    assert entry.label_text == "图 1-1"


def test_registry_lookup_missing_returns_none():
    from docx_formatter.references import LabelRegistry

    registry = LabelRegistry()
    assert registry.lookup("fig:missing") is None


def test_wrap_in_bookmark_creates_bookmark_elements():
    from docx_formatter.references import LabelEntry, wrap_in_bookmark

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("图 1-1")

    entry = LabelEntry(
        kind="figure", key="fig:arch", label_text="图 1-1",
        paragraph_index=0, bookmark_name="_fig_arch", bookmark_id=1,
    )
    wrap_in_bookmark(p._p, run._r, entry)

    xml = p._p.xml
    assert "_fig_arch" in xml
    assert "w:bookmarkEnd" in xml


def test_wrap_in_bookmark_preserves_run_content():
    from docx_formatter.references import LabelEntry, wrap_in_bookmark

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("(1-1)")

    entry = LabelEntry(
        kind="formula", key="formula:e", label_text="(1-1)",
        paragraph_index=0, bookmark_name="_formula_e", bookmark_id=2,
    )
    wrap_in_bookmark(p._p, run._r, entry)

    assert run.text == "(1-1)"


def test_make_ref_field_creates_valid_field_xml():
    from docx_formatter.references import make_ref_field

    fragments = make_ref_field("_fig_arch", "图 1-1")

    assert len(fragments) == 5

    # begin: w:r containing w:fldChar begin
    begin_r = fragments[0]
    begin_char = begin_r.find(qn("w:fldChar"))
    assert begin_char is not None
    assert begin_char.get(qn("w:fldCharType")) == "begin"

    # instrText: w:r containing w:instrText
    instr_r = fragments[1]
    instr_text = instr_r.find(qn("w:instrText"))
    assert instr_text is not None
    assert "_fig_arch" in instr_text.text

    # display text: w:r containing w:t
    display_r = fragments[3]
    t = display_r.find(qn("w:t"))
    assert t is not None
    assert t.text == "图 1-1"

    # end: w:r containing w:fldChar end
    end_r = fragments[4]
    end_char = end_r.find(qn("w:fldChar"))
    assert end_char is not None
    assert end_char.get(qn("w:fldCharType")) == "end"


def test_make_ref_field_superscript():
    """Superscript REF field has vertAlign on display run."""
    from docx_formatter.references import make_ref_field

    fragments = make_ref_field("_ref_smith", "[1]", superscript=True)
    display_r = fragments[3]
    rPr = display_r.find(qn("w:rPr"))
    assert rPr is not None
    vertAlign = rPr.find(qn("w:vertAlign"))
    assert vertAlign is not None
    assert vertAlign.get(qn("w:val")) == "superscript"


def test_make_ref_field_no_superscript_by_default():
    """Without superscript, display run has no vertAlign."""
    from docx_formatter.references import make_ref_field

    fragments = make_ref_field("_fig_arch", "图 1-1")
    display_r = fragments[3]
    rPr = display_r.find(qn("w:rPr"))
    assert rPr is None


def test_references_config_loaded_from_template():
    from docx_formatter.config import load_template

    config = load_template(Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"))
    assert hasattr(config, "references")
    assert config.references.title == "参考文献"
    assert config.references.font_cn == "宋体"
    assert config.references.font_en == "Times New Roman"
    assert config.references.number_style == "bracket"


def test_parse_placeholders_finds_single_key():
    from docx_formatter.references import parse_placeholders
    segments = parse_placeholders("如图{fig:arch}所示")
    assert segments == [("text", "如图"), ("ref", "fig:arch"), ("text", "所示")]


def test_parse_placeholders_finds_multiple_keys():
    from docx_formatter.references import parse_placeholders
    segments = parse_placeholders("见{fig:a}和{tab:b}")
    assert segments == [("text", "见"), ("ref", "fig:a"), ("text", "和"), ("ref", "tab:b")]


def test_parse_placeholders_no_keys():
    from docx_formatter.references import parse_placeholders
    segments = parse_placeholders("普通文本")
    assert segments == [("text", "普通文本")]


def test_parse_placeholders_key_at_start():
    from docx_formatter.references import parse_placeholders
    segments = parse_placeholders("{formula:e}是基本方程")
    assert segments == [("ref", "formula:e"), ("text", "是基本方程")]


def test_parse_placeholders_key_at_end():
    from docx_formatter.references import parse_placeholders
    segments = parse_placeholders("根据研究{ref:smith}")
    assert segments == [("text", "根据研究"), ("ref", "ref:smith")]


def _make_figures_config():
    from unittest.mock import MagicMock
    config = MagicMock()
    config.caption_font_cn = "宋体"
    config.caption_font_en = "Times New Roman"
    config.caption_size = "五号"
    config.caption_bold = False
    config.figure_caption_position = "below"
    config.table_caption_position = "above"
    config.numbering = "chapter"
    config.table_style = "three_line"
    return config


def test_auto_number_figure_captions_chapter():
    from docx_formatter.references import LabelRegistry
    from docx_formatter.figures import auto_number_captions

    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("系统架构")
    doc.add_paragraph("流程图")
    doc.add_heading("Chapter 2", level=1)
    doc.add_paragraph("数据图")

    registry = LabelRegistry()
    figure_items = [
        (1, {"key": "fig:arch", "caption": "系统架构"}),
        (2, {"key": "fig:flow", "caption": "流程图"}),
        (4, {"key": "fig:data", "caption": "数据图"}),
    ]
    auto_number_captions(doc, registry, "figure", figure_items, _make_figures_config())

    assert doc.paragraphs[1].text == "图 1-1 系统架构"
    assert doc.paragraphs[2].text == "图 1-2 流程图"
    assert doc.paragraphs[4].text == "图 2-1 数据图"

    assert registry.lookup("fig:arch").label_text == "图 1-1"
    assert registry.lookup("fig:flow").label_text == "图 1-2"
    assert registry.lookup("fig:data").label_text == "图 2-1"


def test_auto_number_figure_captions_continuous():
    from docx_formatter.references import LabelRegistry
    from docx_formatter.figures import auto_number_captions

    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("系统架构")
    doc.add_heading("Chapter 2", level=1)
    doc.add_paragraph("流程图")

    config = _make_figures_config()
    config.numbering = "continuous"
    registry = LabelRegistry()
    figure_items = [
        (1, {"key": "fig:arch", "caption": "系统架构"}),
        (3, {"key": "fig:flow", "caption": "流程图"}),
    ]
    auto_number_captions(doc, registry, "figure", figure_items, config)

    assert doc.paragraphs[1].text == "图 1 系统架构"
    assert doc.paragraphs[3].text == "图 2 流程图"


def test_auto_number_table_captions():
    from docx_formatter.references import LabelRegistry
    from docx_formatter.figures import auto_number_captions

    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("实验数据")

    registry = LabelRegistry()
    table_items = [
        (1, {"key": "tab:data", "caption": "实验数据"}),
    ]
    auto_number_captions(doc, registry, "table", table_items, _make_figures_config())

    assert doc.paragraphs[1].text == "表 1-1 实验数据"
    assert registry.lookup("tab:data").label_text == "表 1-1"


def test_auto_number_creates_bookmarks():
    from docx_formatter.references import LabelRegistry
    from docx_formatter.figures import auto_number_captions

    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("系统架构")

    registry = LabelRegistry()
    figure_items = [
        (1, {"key": "fig:arch", "caption": "系统架构"}),
    ]
    auto_number_captions(doc, registry, "figure", figure_items, _make_figures_config())

    xml = doc.paragraphs[1]._p.xml
    assert "_fig_arch" in xml


def test_generate_bibliography_creates_numbered_entries():
    from docx_formatter.references import LabelRegistry, generate_bibliography

    doc = Document()
    registry = LabelRegistry()
    registry.register("reference", "ref:smith", "[1]", -1)
    registry.register("reference", "ref:zhang", "[2]", -1)

    references = [
        {"key": "ref:smith", "text": "Smith J. Deep Learning. Nature. 2024."},
        {"key": "ref:zhang", "text": "张三. 机器学习导论. 科学出版社. 2023."},
    ]

    from unittest.mock import MagicMock
    config = MagicMock()
    config.title = "参考文献"
    config.title_font_cn = "黑体"
    config.title_font_en = "Times New Roman"
    config.title_size = "小四"
    config.title_bold = True
    config.title_align = "left"
    config.font_cn = "宋体"
    config.font_en = "Times New Roman"
    config.size = "五号"
    config.align = "left"
    config.hanging_indent = "2chars"
    config.number_style = "bracket"

    generate_bibliography(doc, registry, references, config)

    assert doc.paragraphs[0].text == "参考文献"
    assert "Smith J. Deep Learning. Nature. 2024." in doc.paragraphs[1].text
    assert "张三. 机器学习导论. 科学出版社. 2023." in doc.paragraphs[2].text
    assert "_ref_smith" in doc.paragraphs[1]._p.xml
    assert "_ref_zhang" in doc.paragraphs[2]._p.xml


def test_generate_bibliography_updates_registry_indices():
    from docx_formatter.references import LabelRegistry, generate_bibliography

    doc = Document()
    doc.add_paragraph("Some body text")

    registry = LabelRegistry()
    registry.register("reference", "ref:smith", "[1]", -1)

    references = [
        {"key": "ref:smith", "text": "Smith J. Deep Learning. Nature. 2024."},
    ]

    from unittest.mock import MagicMock
    config = MagicMock()
    config.title = "参考文献"
    config.title_font_cn = "黑体"
    config.title_font_en = "Times New Roman"
    config.title_size = "小四"
    config.title_bold = True
    config.title_align = "left"
    config.font_cn = "宋体"
    config.font_en = "Times New Roman"
    config.size = "五号"
    config.align = "left"
    config.hanging_indent = "2chars"
    config.number_style = "bracket"

    generate_bibliography(doc, registry, references, config)

    entry = registry.lookup("ref:smith")
    assert entry.paragraph_index == 2  # 0=body, 1=title, 2=first ref entry


def test_create_document_with_cross_references(tmp_path):
    from docx_formatter.pipeline import create_document

    image_path = Path("D:/2026project/docx_skill/.worktrees/docx-formatter/examples/image.png")
    if not image_path.exists():
        return  # skip if test image not available

    output = create_document(
        output_path=tmp_path / "xref.docx",
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"),
        paragraphs=[
            {"type": "heading_1", "text": "第一章 实验"},
            {"type": "body", "text": "如{fig:arch}所示，系统架构如下。"},
            {"type": "image", "path": str(image_path),
             "key": "fig:arch", "caption": "系统架构"},
            {"type": "formula", "latex": "E=mc^2", "key": "formula:emc2"},
            {"type": "body", "text": "由公式{formula:emc2}可得。"},
        ],
        overrides={},
    )

    doc = Document(str(output))
    # Body paragraph with REF field
    body_xml = doc.paragraphs[1]._p.xml
    assert "_fig_arch" in body_xml
    assert "REF" in body_xml

    # Formula reference
    body_xml2 = doc.paragraphs[5]._p.xml
    assert "_formula_emc2" in body_xml2
    assert "REF" in body_xml2


def test_create_document_with_bibliography(tmp_path):
    from docx_formatter.pipeline import create_document

    output = create_document(
        output_path=tmp_path / "bib.docx",
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"),
        paragraphs=[
            {"type": "heading_1", "text": "第一章"},
            {"type": "body", "text": "根据研究{ref:smith}，效果显著。"},
        ],
        overrides={},
        references=[
            {"key": "ref:smith", "text": "Smith J. Deep Learning. Nature. 2024."},
        ],
    )

    doc = Document(str(output))
    texts = [p.text for p in doc.paragraphs]
    assert "参考文献" in texts
    assert any("Smith J. Deep Learning" in t for t in texts)
    assert any("[1]" in t for t in texts)

    # REF field in body
    body_xml = doc.paragraphs[1]._p.xml
    assert "_ref_smith" in body_xml


def test_format_document_preserves_bookmarks(tmp_path):
    """format_document preserves bookmarks created by create_document."""
    from docx_formatter.pipeline import format_document, create_document

    source = tmp_path / "source.docx"
    create_document(
        output_path=source,
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"),
        paragraphs=[
            {"type": "heading_1", "text": "第一章"},
            {"type": "formula", "latex": "x+y", "key": "formula:sum"},
        ],
        overrides={},
    )

    result = format_document(
        source_path=source,
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"),
        overrides={},
    )

    doc = Document(str(result))
    formula_p = doc.paragraphs[1]
    assert "_formula_sum" in formula_p._p.xml


def test_format_document_handles_figure_captions(tmp_path):
    """format_document still formats figure captions correctly."""
    from docx_formatter.pipeline import format_document

    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("图 1-1 系统架构")
    doc.add_paragraph("图 1-2 流程图")
    doc.save(source)

    result = format_document(
        source_path=source,
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"),
        overrides={},
    )

    formatted = Document(str(result))
    assert "图 1-1" in formatted.paragraphs[1].text
    assert formatted.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
