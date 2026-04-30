from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def test_extract_format_from_styled_docx(tmp_path):
    from docx_formatter.format_extractor import extract_format_from_docx

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    heading_style = doc.styles["Heading 1"]
    heading_style.font.name = "Times New Roman"
    heading_style._element.get_or_add_rPr().rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "黑体"
    )
    heading_style.font.size = 16 * 12700
    heading_style.font.bold = True
    heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.get_or_add_rPr().rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体"
    )
    normal_style.font.size = 12 * 12700
    normal_style.paragraph_format.first_line_indent = 24 * 12700
    normal_style.paragraph_format.line_spacing = 1.5

    doc.sections[0].header.paragraphs[0].text = "毕业论文"

    path = tmp_path / "styled.docx"
    doc.save(path)

    config = extract_format_from_docx(path)

    assert config.page.size == "A4"
    assert config.page.margins.top == "2.54cm"
    assert config.page.margins.left == "3.17cm"
    assert config.page.header.content == "毕业论文"
    assert config.headings.level_1.font_cn == "黑体"
    assert config.headings.level_1.font_en == "Times New Roman"
    assert config.headings.level_1.bold is True
    assert config.headings.level_1.align == "center"
    assert config.body.font_cn == "宋体"
    assert config.body.font_en == "Times New Roman"
    assert config.body.first_line_indent == "2chars"
    assert config.body.line_spacing == "1.5lines"


def test_extract_format_from_text_spec_docx(tmp_path):
    from docx_formatter.format_extractor import extract_format_from_docx

    doc = Document()
    doc.add_paragraph("标题（一）用黑体三号居中")
    doc.add_paragraph("正文用宋体小四，首行缩进2字符，行距1.5倍")
    doc.add_paragraph("页边距：上下2.54cm，左右3.17cm")
    doc.add_paragraph("页眉内容：毕业论文")
    path = tmp_path / "spec.docx"
    doc.save(path)

    config = extract_format_from_docx(path)

    assert config.headings.level_1.font_cn == "黑体"
    assert config.headings.level_1.size == "三号"
    assert config.headings.level_1.align == "center"
    assert config.body.font_cn == "宋体"
    assert config.body.size == "小四"
    assert config.body.first_line_indent == "2chars"
    assert config.body.line_spacing == "1.5lines"
    assert config.page.margins.top == "2.54cm"
    assert config.page.margins.left == "3.17cm"
    assert config.page.header.content == "毕业论文"


def test_extract_format_prefers_text_spec_over_default_builtin_styles(tmp_path):
    from docx_formatter.format_extractor import extract_format_from_docx

    doc = Document()
    doc.add_paragraph("标题（一）用黑体三号居中")
    path = tmp_path / "prefer_text_spec.docx"
    doc.save(path)

    config = extract_format_from_docx(path)

    assert config.headings.level_1.size == "三号"


def test_extract_format_falls_back_to_text_spec_when_styles_not_informative(tmp_path):
    from docx_formatter.format_extractor import extract_format_from_docx

    doc = Document()
    doc.add_paragraph("标题（一）用黑体三号居中")
    doc.add_paragraph("正文用宋体小四")
    path = tmp_path / "fallback.docx"
    doc.save(path)

    config = extract_format_from_docx(path)

    assert config.headings.level_1.font_cn == "黑体"
    assert config.body.font_cn == "宋体"
