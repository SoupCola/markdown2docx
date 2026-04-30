from pathlib import Path

from docx import Document

from docx_formatter.pipeline import create_document, format_document


def test_format_document_saves_formatted_copy(tmp_path):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_heading("第一章 绪论", level=1)
    document.add_paragraph("正文内容")
    document.save(source)

    output = format_document(
        source_path=source,
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"),
        overrides={},
    )

    assert output.exists()
    assert output.name == "source_formatted.docx"


def test_create_document_builds_new_docx_from_sections(tmp_path):
    output = create_document(
        output_path=tmp_path / "new.docx",
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/report.yaml"),
        paragraphs=[
            {"type": "heading_1", "text": "项目概述"},
            {"type": "body", "text": "这里是正文。"},
        ],
        overrides={},
    )

    created = Document(output)
    assert created.paragraphs[0].text == "项目概述"
    assert created.paragraphs[1].text == "这里是正文。"




def test_create_document_builds_word_numbered_paragraphs(tmp_path):
    output = create_document(
        output_path=tmp_path / "numbered.docx",
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"),
        paragraphs=[
            {"type": "heading_1", "text": "第一章 绪论"},
            {"type": "numbered_level_1", "text": "研究目标"},
            {"type": "numbered_level_2", "text": "理论基础"},
            {"type": "numbered_bracket", "text": "参考文献1"},
        ],
        overrides={},
    )

    created = Document(output)
    assert created.paragraphs[1].text == "研究目标"
    assert "w:numPr" in created.paragraphs[1]._p.xml
    assert "w:numPr" in created.paragraphs[2]._p.xml
    assert "w:numPr" in created.paragraphs[3]._p.xml


def test_create_document_supports_heading_levels_without_body_indent(tmp_path):
    output = create_document(
        output_path=tmp_path / "levels.docx",
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"),
        paragraphs=[
            {"type": "heading_1", "text": "第一章 绪论"},
            {"type": "heading_2", "text": "1.1 研究背景"},
            {"type": "heading_3", "text": "1.1.1 研究目标"},
            {"type": "body", "text": "这里是正文。"},
        ],
        overrides={},
    )

    created = Document(output)
    assert created.paragraphs[1].style.name == "Heading 2"
    assert created.paragraphs[2].style.name == "Heading 3"
    assert created.paragraphs[1].paragraph_format.first_line_indent is None
    assert created.paragraphs[2].paragraph_format.first_line_indent is None


def test_format_document_applies_heading_body_and_page_changes(tmp_path):
    source = tmp_path / "integration.docx"
    document = Document()
    document.add_heading("第一章 绪论", level=1)
    document.add_paragraph("这是第一段正文。")
    document.add_paragraph("图 1-1 系统架构")
    document.save(source)

    output = format_document(
        source_path=source,
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"),
        overrides={"body": {"first_line_indent": "2chars"}},
    )

    formatted = Document(output)
    assert formatted.paragraphs[0].style.name == "Heading 1"
    assert round(formatted.paragraphs[1].paragraph_format.first_line_indent.pt, 1) == 24.0
    assert "PAGE" in formatted.sections[0].footer._element.xml


def test_skill_md_mentions_confirmation_and_safe_output_path():
    content = Path("D:/2026project/docx_skill/.worktrees/docx-formatter/skill.md").read_text(encoding="utf-8")
    assert "确认" in content
    assert "_formatted.docx" in content
    assert "模板" in content


THESIS_TEMPLATE = Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml")


def test_format_document_converts_supported_numbering_and_preserves_ambiguous_text(tmp_path):
    source = tmp_path / "numbering-source.docx"
    document = Document()
    document.add_paragraph("[1] 参考文献1")
    document.add_paragraph("[2] 参考文献2")
    document.add_paragraph("2. 不安全编号")
    document.save(source)

    output = format_document(
        source_path=source,
        template_path=THESIS_TEMPLATE,
        overrides={},
    )

    formatted = Document(output)
    assert formatted.paragraphs[0].text == "参考文献1"
    assert formatted.paragraphs[1].text == "参考文献2"
    assert "w:numPr" in formatted.paragraphs[0]._p.xml
    assert formatted.paragraphs[2].text == "2. 不安全编号"
    assert "w:numPr" not in formatted.paragraphs[2]._p.xml


def test_format_document_keeps_orphan_level_2_text_unchanged(tmp_path):
    source = tmp_path / "orphan-level-2.docx"
    document = Document()
    document.add_paragraph("（2）孤立二级项")
    document.save(source)

    output = format_document(
        source_path=source,
        template_path=THESIS_TEMPLATE,
        overrides={},
    )

    formatted = Document(output)
    assert formatted.paragraphs[0].text == "（2）孤立二级项"
    assert "w:numPr" not in formatted.paragraphs[0]._p.xml


def test_generated_numbering_preserves_body_text_without_prefixes(tmp_path):
    output = create_document(
        output_path=tmp_path / "generated-numbering.docx",
        template_path=Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"),
        paragraphs=[
            {"type": "numbered_level_1", "text": "第一项"},
            {"type": "numbered_level_1", "text": "第二项"},
            {"type": "numbered_bracket", "text": "参考文献1"},
            {"type": "numbered_bracket", "text": "参考文献2"},
        ],
        overrides={},
    )

    created = Document(output)
    assert [p.text for p in created.paragraphs] == ["第一项", "第二项", "参考文献1", "参考文献2"]
    assert all("w:numPr" in p._p.xml for p in created.paragraphs)


def _get_numId(paragraph):
    from docx.oxml.ns import qn
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    return numPr.find(qn("w:numId")).get(qn("w:val"))


def test_create_document_restarts_numbering_across_groups(tmp_path):
    """Numbered items separated by non-numbered content should restart from 1."""
    output = create_document(
        output_path=tmp_path / "groups.docx",
        template_path=THESIS_TEMPLATE,
        paragraphs=[
            {"type": "numbered_level_1", "text": "第一组第一项"},
            {"type": "numbered_level_1", "text": "第一组第二项"},
            {"type": "body", "text": "正文分隔"},
            {"type": "numbered_level_1", "text": "第二组第一项"},
            {"type": "numbered_level_2", "text": "第二组二级"},
        ],
        overrides={},
    )

    created = Document(output)
    assert _get_numId(created.paragraphs[0]) != _get_numId(created.paragraphs[3])


def test_create_document_numbered_items_have_no_first_line_indent(tmp_path):
    """Numbered items should not inherit body first-line indent."""
    output = create_document(
        output_path=tmp_path / "no-indent.docx",
        template_path=THESIS_TEMPLATE,
        paragraphs=[
            {"type": "body", "text": "正文有缩进"},
            {"type": "numbered_level_1", "text": "编号无缩进"},
            {"type": "numbered_bracket", "text": "参考文献无缩进"},
        ],
        overrides={},
    )

    created = Document(output)
    body_indent = created.paragraphs[0].paragraph_format.first_line_indent
    assert body_indent is not None and body_indent.pt > 0
    numbered_indent_1 = created.paragraphs[1].paragraph_format.first_line_indent
    numbered_indent_2 = created.paragraphs[2].paragraph_format.first_line_indent
    assert numbered_indent_1 is None or numbered_indent_1 == 0
    assert numbered_indent_2 is None or numbered_indent_2 == 0


def test_format_document_clears_indent_on_converted_numbering(tmp_path):
    """Converted numbering items should not retain body first-line indent."""
    source = tmp_path / "bracket-indent.docx"
    document = Document()
    document.add_paragraph("[1] 参考文献一")
    document.add_paragraph("[2] 参考文献二")
    document.save(source)

    output = format_document(
        source_path=source,
        template_path=THESIS_TEMPLATE,
        overrides={},
    )

    formatted = Document(output)
    indent_0 = formatted.paragraphs[0].paragraph_format.first_line_indent
    indent_1 = formatted.paragraphs[1].paragraph_format.first_line_indent
    assert indent_0 is None or indent_0 == 0
    assert indent_1 is None or indent_1 == 0


def test_create_document_with_formula(tmp_path):
    """create_document supports formula type with LaTeX."""
    from docx_formatter.formulas import is_formula_paragraph

    output = create_document(
        output_path=tmp_path / "formula.docx",
        template_path=THESIS_TEMPLATE,
        paragraphs=[
            {"type": "heading_1", "text": "第一章"},
            {"type": "formula", "latex": "E=mc^2"},
            {"type": "body", "text": "这是一段正文"},
        ],
        overrides={},
    )

    doc = Document(str(output))
    formula_found = any(is_formula_paragraph(p) for p in doc.paragraphs)
    assert formula_found


def test_format_document_preserves_formulas(tmp_path):
    """format_document preserves formula content and applies formatting."""
    from docx_formatter.formulas import is_formula_paragraph

    # First create a doc with formula
    source = tmp_path / "source.docx"
    create_document(
        output_path=source,
        template_path=THESIS_TEMPLATE,
        paragraphs=[
            {"type": "heading_1", "text": "第一章"},
            {"type": "formula", "latex": "x^2 + y^2 = z^2"},
            {"type": "body", "text": "这是正文"},
        ],
        overrides={},
    )

    result = format_document(
        source_path=source,
        template_path=THESIS_TEMPLATE,
        overrides={},
    )

    doc = Document(str(result))
    formula_found = any(is_formula_paragraph(p) for p in doc.paragraphs)
    assert formula_found


def test_create_document_formula_has_tab_stops(tmp_path):
    """Formula paragraphs have center and right tab stops for alignment."""
    from docx_formatter.formulas import is_formula_paragraph
    from docx.oxml.ns import qn

    output = create_document(
        output_path=tmp_path / "tabs.docx",
        template_path=THESIS_TEMPLATE,
        paragraphs=[
            {"type": "heading_1", "text": "第一章"},
            {"type": "formula", "latex": "E=mc^2"},
        ],
        overrides={},
    )

    doc = Document(str(output))
    for p in doc.paragraphs:
        if is_formula_paragraph(p):
            pPr = p._p.find(qn("w:pPr"))
            tabs = pPr.find(qn("w:tabs")) if pPr is not None else None
            assert tabs is not None
            tab_elements = tabs.findall(qn("w:tab"))
            vals = [t.get(qn("w:val")) for t in tab_elements]
            assert "center" in vals
            assert "right" in vals
            break
    else:
        assert False, "No formula paragraph found"


def test_create_document_with_image(tmp_path):
    """create_document supports image insertion from local path."""
    image_path = Path("D:/2026project/docx_skill/.worktrees/docx-formatter/examples/image.png")
    if not image_path.exists():
        return  # skip if test image not available

    output = create_document(
        output_path=tmp_path / "image.docx",
        template_path=THESIS_TEMPLATE,
        paragraphs=[
            {"type": "heading_1", "text": "第一章 实验结果"},
            {"type": "body", "text": "实验结果如下图所示。"},
            {"type": "image", "path": str(image_path), "caption": "图 1-1 实验结果"},
        ],
        overrides={},
    )

    doc = Document(str(output))
    # Check that the document was created successfully
    assert output.exists()
    # Image paragraph should be centered
    # The image is in paragraph at index 2, caption at index 3
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    assert doc.paragraphs[2].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert doc.paragraphs[3].text == "图 1-1 实验结果"
    assert doc.paragraphs[3].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
