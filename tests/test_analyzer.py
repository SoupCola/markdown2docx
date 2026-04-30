from docx import Document

from docx_formatter.analyzer import analyze, summarize_analysis


def test_analyze_reports_headings_paragraphs_tables_and_sections(tmp_path):
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("第一章 绪论", level=1)
    document.add_paragraph("这是正文")
    document.add_table(rows=2, cols=2)
    document.save(source)

    result = analyze(source)

    assert result["total_paragraphs"] == 2
    assert result["headings_count"]["level_1"] == 1
    assert len(result["tables"]) == 1
    assert result["sections"][0]["page_size"]


def test_summarize_analysis_returns_user_readable_text(tmp_path):
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("标题", level=1)
    document.add_paragraph("正文")
    document.save(source)

    summary = summarize_analysis(analyze(source))
    assert "段落数" in summary
    assert "一级标题" in summary
