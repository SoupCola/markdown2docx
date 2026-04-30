from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


THESIS_TEMPLATE = Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml")
SKILL_ROOT = Path(__file__).resolve().parents[1]


import sys

sys.path.insert(0, str(SKILL_ROOT))
for _module_name in [name for name in list(sys.modules) if name == "docx_formatter" or name.startswith("docx_formatter.")]:
    sys.modules.pop(_module_name, None)

from docx_formatter.pipeline import create_from_markdown


def _document_xml(docx_path: Path) -> str:
    with ZipFile(docx_path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def test_create_from_markdown_builds_docx_from_template_and_markdown(tmp_path):
    md_path = tmp_path / "chapter1.md"
    md_path.write_text("# 第一章\n\n正文第一段。\n", encoding="utf-8")

    output = create_from_markdown(
        output_path=tmp_path / "from-markdown.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
    )

    created = Document(output)
    assert created.paragraphs[0].text == "第一章"
    assert created.paragraphs[1].text == "正文第一段。"


def test_create_from_markdown_builds_figure_and_table_ref_fields(tmp_path):
    md_path = tmp_path / "figure-table-xrefs.md"
    images_dir = tmp_path / "images"
    images_dir.mkdir()

    source = Document()
    source.add_paragraph("test image")
    image_docx = images_dir / "overall-design.docx"
    source.save(image_docx)

    md_path.write_text(
        "系统总体结构如图 4-1 所示。\n\n"
        "![图4-1 系统总体设计结构图](images/overall-design.png)\n\n"
        "各层职责如表 4-1 所示。\n\n"
        "| 层次 | 组成内容 | 主要职责 |\n"
        "| --- | --- | --- |\n"
        "| 用户层 | 普通用户 | 发起处理任务 |\n",
        encoding="utf-8",
    )

    from PIL import Image

    image_path = images_dir / "overall-design.png"
    Image.new("RGB", (10, 10), color="white").save(image_path)

    output = create_from_markdown(
        output_path=tmp_path / "figure-table-xrefs.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
    )

    created = Document(output)
    assert "{fig:auto_1}" not in created.paragraphs[0].text
    assert "{tab:auto_1}" not in created.paragraphs[2].text

    document_xml = _document_xml(output)
    assert document_xml.count(" REF _fig_auto_1 \\h ") == 1
    assert document_xml.count(" REF _tab_auto_1 \\h ") == 1
    assert 'w:name="_fig_auto_1"' in document_xml
    assert 'w:name="_tab_auto_1"' in document_xml


def test_create_from_markdown_registers_table_ref_field_without_explicit_table_caption(tmp_path):
    md_path = tmp_path / "table-xref-no-caption.md"
    md_path.write_text(
        "各层职责如表 4-1 所示。\n\n"
        "| 层次 | 组成内容 | 主要职责 |\n"
        "| --- | --- | --- |\n"
        "| 用户层 | 普通用户 | 发起处理任务 |\n",
        encoding="utf-8",
    )

    output = create_from_markdown(
        output_path=tmp_path / "table-xref-no-caption.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
    )

    created = Document(output)
    assert "{tab:auto_1}" not in created.paragraphs[0].text

    document_xml = _document_xml(output)
    assert document_xml.count(" REF _tab_auto_1 \\h ") == 1
    assert 'w:name="_tab_auto_1"' in document_xml


def test_create_from_markdown_uses_docx_format_as_overrides(tmp_path):
    format_docx = tmp_path / "format.docx"
    source = Document()
    section = source.sections[0]
    section.top_margin = Cm(2.0)

    heading_style = source.styles["Heading 1"]
    heading_style.font.name = "Times New Roman"
    heading_style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    normal_style = source.styles["Normal"]
    rpr = normal_style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.first_line_indent = Pt(24)
    normal_style.paragraph_format.line_spacing = 1.5
    source.save(format_docx)

    md_path = tmp_path / "chapter1.md"
    md_path.write_text("# 第一章\n\n正文第一段。\n", encoding="utf-8")

    output = create_from_markdown(
        output_path=tmp_path / "from-format-docx.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
        format_docx_path=format_docx,
        overrides={
            "headings": {
                "level_2": {"size": "小三"},
                "level_3": {"size": "四号"},
            }
        },
    )

    created = Document(output)
    assert round(created.sections[0].top_margin.cm, 2) == 2.0
    assert created.paragraphs[0].style.name == "Heading 1"
    assert round(created.paragraphs[1].paragraph_format.first_line_indent.pt, 1) == 24.0
    assert created.paragraphs[1].text == "正文第一段。"


def test_create_from_markdown_builds_bibliography_and_ref_field_for_footnote_citation(tmp_path):
    md_path = tmp_path / "refs.md"
    md_path.write_text(
        "根据研究[^smith2024]，效果显著。\n\n[^smith2024]: Smith J. Deep Learning. Nature, 2024.\n",
        encoding="utf-8",
    )

    output = create_from_markdown(
        output_path=tmp_path / "refs.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
    )

    created = Document(output)
    assert "{ref:smith2024}" not in created.paragraphs[0].text
    assert "smith2024" not in created.paragraphs[0].text
    assert created.paragraphs[-2].text == "参考文献"
    assert created.paragraphs[-1].text == "[1] Smith J. Deep Learning. Nature, 2024."



def test_create_from_markdown_builds_bibliography_and_ref_field_for_numbered_citation(tmp_path):
    md_path = tmp_path / "numbered-refs.md"
    md_path.write_text(
        "根据研究[1]，效果显著。\n\n[1] Smith J. Deep Learning. Nature, 2024.\n",
        encoding="utf-8",
    )

    output = create_from_markdown(
        output_path=tmp_path / "numbered-refs.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
    )

    created = Document(output)
    assert "{ref:1}" not in created.paragraphs[0].text
    assert created.paragraphs[-2].text == "参考文献"
    assert created.paragraphs[-1].text == "[1] Smith J. Deep Learning. Nature, 2024."



def test_create_from_markdown_builds_multiple_ref_fields_in_one_paragraph(tmp_path):
    md_path = tmp_path / "multi-refs.md"
    md_path.write_text(
        "根据研究[^smith2024]与[2]，效果显著。\n\n[^smith2024]: Smith J. Deep Learning. Nature, 2024.\n[2] Wang L. Document Intelligence. Science, 2025.\n",
        encoding="utf-8",
    )

    output = create_from_markdown(
        output_path=tmp_path / "multi-refs.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
    )

    created = Document(output)
    assert "{ref:smith2024}" not in created.paragraphs[0].text
    assert "{ref:2}" not in created.paragraphs[0].text
    assert created.paragraphs[1].text == "参考文献"
    assert created.paragraphs[2].text == "[1] Smith J. Deep Learning. Nature, 2024."
    assert created.paragraphs[3].text == "[2] Wang L. Document Intelligence. Science, 2025."

    document_xml = _document_xml(output)
    assert document_xml.count(" REF _ref_smith2024 \\h ") == 1
    assert document_xml.count(" REF _ref_2 \\h ") == 1
    assert 'w:name="_ref_smith2024"' in document_xml
    assert 'w:name="_ref_2"' in document_xml



def test_create_from_markdown_numbers_references_by_first_citation_order(tmp_path):
    md_path = tmp_path / "citation-order.md"
    md_path.write_text(
        "先引用[2]，再引用[1]。\n\n"
        "[1] Smith J. Deep Learning. Nature, 2024.\n"
        "[2] Wang L. Document Intelligence. Science, 2025.\n",
        encoding="utf-8",
    )

    output = create_from_markdown(
        output_path=tmp_path / "citation-order.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
    )

    created = Document(output)
    assert created.paragraphs[1].text == "参考文献"
    assert created.paragraphs[2].text == "[1] Wang L. Document Intelligence. Science, 2025."
    assert created.paragraphs[3].text == "[2] Smith J. Deep Learning. Nature, 2024."

    document_xml = _document_xml(output)
    assert document_xml.count(" REF _ref_2 \\h ") == 1
    assert document_xml.count(" REF _ref_1 \\h ") == 1



def test_create_from_markdown_ignores_manual_bibliography_heading(tmp_path):
    md_path = tmp_path / "manual-bibliography-heading.md"
    md_path.write_text(
        "正文引用[1]。\n\n"
        "# 参考文献\n\n"
        "[1] Smith J. Deep Learning. Nature, 2024.\n",
        encoding="utf-8",
    )

    output = create_from_markdown(
        output_path=tmp_path / "manual-bibliography-heading.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
    )

    created = Document(output)
    assert [paragraph.text for paragraph in created.paragraphs].count("参考文献") == 1
    assert created.paragraphs[1].text == "参考文献"
    assert created.paragraphs[2].text == "[1] Smith J. Deep Learning. Nature, 2024."


def test_create_from_markdown_uses_local_skill_runtime_for_tables(tmp_path):
    import importlib
    import sys

    sys.path.insert(0, str(SKILL_ROOT))
    sys.modules.pop("docx_formatter", None)
    for name in list(sys.modules):
        if name.startswith("docx_formatter."):
            sys.modules.pop(name, None)

    pipeline = importlib.import_module("docx_formatter.pipeline")

    md_path = tmp_path / "tables-runtime.md"
    md_path.write_text(
        "# 运行时\n\n"
        "| 列1 | 列2 |\n"
        "| --- | --- |\n"
        "| A | B |\n",
        encoding="utf-8",
    )

    output = pipeline.create_from_markdown(
        output_path=tmp_path / "tables-runtime.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
    )

    created = Document(output)
    assert len(created.tables) == 1


def test_create_from_markdown_renders_markdown_tables_into_docx_tables(tmp_path):
    md_path = tmp_path / "tables.md"
    md_path.write_text(
        "# 测试环境\n\n"
        "| 项目 | 配置情况 | 说明 |\n"
        "| --- | --- | --- |\n"
        "| GPU | RTX 3060 | 用于推理 |\n"
        "| 操作系统 | Windows 11 | 运行环境 |\n",
        encoding="utf-8",
    )

    output = create_from_markdown(
        output_path=tmp_path / "tables.docx",
        md_paths=[md_path],
        template_path=THESIS_TEMPLATE,
    )

    created = Document(output)
    assert len(created.tables) == 1
    assert created.tables[0].cell(0, 0).text == "项目"
    assert created.tables[0].cell(0, 1).text == "配置情况"
    assert created.tables[0].cell(0, 2).text == "说明"
    assert created.tables[0].cell(1, 0).text == "GPU"
    assert created.tables[0].cell(1, 1).text == "RTX 3060"
    assert created.tables[0].cell(1, 2).text == "用于推理"
    assert created.tables[0].cell(2, 0).text == "操作系统"
    assert created.tables[0].cell(2, 1).text == "Windows 11"
    assert created.tables[0].cell(2, 2).text == "运行环境"
