from pathlib import Path

from docx import Document

from docx_formatter.config import load_template
from docx_formatter.numbering import apply_generated_numbering, convert_detected_numbering, detect_numbering_groups


THESIS_TEMPLATE = Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml")


def test_detect_numbering_groups_recognizes_supported_prefixes():
    document = Document()
    document.add_paragraph("1. 一级内容")
    document.add_paragraph("（1）二级内容")
    document.add_paragraph("[1] 参考文献1")
    config = load_template(THESIS_TEMPLATE)

    analysis = detect_numbering_groups(document, config.numbering)

    assert any(item.text == "一级内容" for item in analysis.convertible_items)
    assert any(item.text == "二级内容" for item in analysis.convertible_items)
    assert any(item.text == "参考文献1" for item in analysis.convertible_items)


def test_detect_numbering_groups_rejects_non_starting_or_non_contiguous_sequences():
    document = Document()
    document.add_paragraph("2. 内容")
    document.add_paragraph("[2] 参考文献2")
    config = load_template(THESIS_TEMPLATE)

    analysis = detect_numbering_groups(document, config.numbering)

    reasons = {item.reason for item in analysis.rejected_items}
    assert "sequence_must_start_at_1" in reasons


def test_detect_numbering_groups_rejects_figure_like_prefixes():
    document = Document()
    document.add_paragraph("表 1. 测试数据")
    document.add_paragraph("图 1. 系统架构")
    config = load_template(THESIS_TEMPLATE)

    analysis = detect_numbering_groups(document, config.numbering)

    assert analysis.convertible_items == []
    assert analysis.rejected_items == []


def test_detect_numbering_groups_assigns_level_2_items_under_level_1_parent():
    document = Document()
    document.add_paragraph("1. 一级内容")
    document.add_paragraph("（1）二级内容")
    document.add_paragraph("（2）二级补充")
    document.add_paragraph("2. 第二项")
    config = load_template(THESIS_TEMPLATE)

    analysis = detect_numbering_groups(document, config.numbering)

    level_2_items = [item for item in analysis.convertible_items if item.kind == "level_2"]
    assert [item.text for item in level_2_items] == ["二级内容", "二级补充"]


def test_detect_numbering_groups_rejects_orphan_level_2_items():
    document = Document()
    document.add_paragraph("（2）脱离父级")
    config = load_template(THESIS_TEMPLATE)

    analysis = detect_numbering_groups(document, config.numbering)

    assert analysis.rejected_items[0].reason == "missing_parent_level_1_context"


def test_apply_generated_numbering_assigns_numbering_xml_to_supported_items():
    document = Document()
    paragraphs = [
        {"type": "numbered_level_1", "text": "研究目标"},
        {"type": "numbered_level_2", "text": "理论基础"},
        {"type": "numbered_bracket", "text": "参考文献1"},
    ]

    apply_generated_numbering(document, paragraphs, load_template(THESIS_TEMPLATE).numbering)

    assert document.paragraphs[0].text == "研究目标"
    assert "w:numPr" in document.paragraphs[0]._p.xml
    assert "w:ilvl" in document.paragraphs[1]._p.xml
    assert "w:numPr" in document.paragraphs[2]._p.xml


def test_convert_detected_numbering_replaces_bracket_prefix_with_word_numbering():
    document = Document()
    document.add_paragraph("[1] 参考文献1")
    document.add_paragraph("[2] 参考文献2")
    config = load_template(THESIS_TEMPLATE)

    analysis = detect_numbering_groups(document, config.numbering)
    convert_detected_numbering(document, analysis, config.numbering)

    assert document.paragraphs[0].text == "参考文献1"
    assert document.paragraphs[1].text == "参考文献2"
    assert "w:numPr" in document.paragraphs[0]._p.xml
    assert "w:numPr" in document.paragraphs[1]._p.xml


def test_convert_detected_numbering_handles_bracket_prefix_without_space():
    document = Document()
    document.add_paragraph("[1]文献1")
    config = load_template(THESIS_TEMPLATE)

    analysis = detect_numbering_groups(document, config.numbering)
    convert_detected_numbering(document, analysis, config.numbering)

    assert document.paragraphs[0].text == "文献1"
    assert "w:numPr" in document.paragraphs[0]._p.xml


def test_detect_numbering_groups_rejects_skipped_level_1_sequence():
    document = Document()
    document.add_paragraph("1. 第一项")
    document.add_paragraph("3. 第三项")
    config = load_template(THESIS_TEMPLATE)

    analysis = detect_numbering_groups(document, config.numbering)

    assert {item.reason for item in analysis.rejected_items} == {"sequence_must_be_contiguous"}


def test_detect_numbering_groups_splits_sequences_at_paragraph_gaps():
    """Level 1 items in different chapters should be detected as separate sequences."""
    document = Document()
    document.add_paragraph("1. 第一章第一项")
    document.add_paragraph("2. 第一章第二项")
    document.add_paragraph("（1）二级内容")
    document.add_paragraph("3. 第一章第三项")
    document.add_paragraph("正文分隔")
    document.add_paragraph("1. 第二章第一项")
    document.add_paragraph("2. 第二章第二项")
    config = load_template(THESIS_TEMPLATE)

    analysis = detect_numbering_groups(document, config.numbering)

    level_1_texts = [item.text for item in analysis.convertible_items if item.kind == "level_1"]
    assert level_1_texts == [
        "第一章第一项", "第一章第二项", "第一章第三项",
        "第二章第一项", "第二章第二项",
    ]
    level_2_texts = [item.text for item in analysis.convertible_items if item.kind == "level_2"]
    assert level_2_texts == ["二级内容"]
