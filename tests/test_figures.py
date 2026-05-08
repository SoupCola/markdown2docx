from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_formatter.config import load_template
from docx_formatter.figures import apply_figure_table_formatting


def test_apply_figure_table_formatting_formats_caption_paragraphs():
    document = Document()
    document.add_paragraph("图 1-1 系统架构")
    document.add_paragraph("表 1-1 测试数据")
    table = document.add_table(rows=2, cols=2)
    config = load_template(Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"))

    apply_figure_table_formatting(document, config.figures)

    figure_caption = document.paragraphs[0]
    table_caption = document.paragraphs[1]
    assert figure_caption.alignment is not None
    assert table_caption.alignment is not None
    assert table._tbl.xml.count("w:top") >= 1


def test_apply_figure_table_formatting_bolds_header_and_centers_cells():
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "字号"
    table.cell(1, 1).text = "小四"
    config = load_template(Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"))

    apply_figure_table_formatting(document, config.figures)

    assert table.cell(0, 0).paragraphs[0].runs[0].bold is True
    assert table.cell(0, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert table.cell(1, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert table.cell(0, 0).vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
    assert table.cell(1, 1).vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
