from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from docx_formatter.config import load_template
from docx_formatter.headings import apply_heading_formatting


def test_apply_heading_formatting_updates_heading_paragraphs():
    document = Document()
    heading = document.add_heading("第一章 绪论", level=1)
    config = load_template(Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"))

    apply_heading_formatting(document, config.headings)

    assert heading.style.name == "Heading 1"
    assert heading.paragraph_format.alignment is not None


def test_apply_heading_formatting_updates_heading_runs_font_and_color():
    document = Document()
    heading = document.add_heading("第一章 绪论", level=1)
    config = load_template(Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates/thesis.yaml"))

    apply_heading_formatting(document, config.headings)

    assert heading.runs[0]._element.rPr.rFonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia") == "黑体"
    assert heading.runs[0].font.color.rgb == RGBColor(0x00, 0x00, 0x00)
