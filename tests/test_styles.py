from dataclasses import replace
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_formatter.config import load_template
from docx_formatter.styles import apply_font, get_or_create_style, register_heading_style


TEMPLATES_DIR = Path("D:/2026project/docx_skill/.worktrees/docx-formatter/templates")


def test_get_or_create_style_returns_same_style_instance():
    document = Document()
    style_a = get_or_create_style(document, "CustomHeading", 1)
    style_b = get_or_create_style(document, "CustomHeading", 1)
    assert style_a.style_id == style_b.style_id


def test_apply_font_sets_run_font_fields():
    document = Document()
    run = document.add_paragraph().add_run("示例")
    apply_font(run, "宋体", "Times New Roman", "小四", True, False)
    assert run.font.name == "Times New Roman"
    assert run.bold is True


def test_register_heading_style_sets_alignment_and_spacing():
    document = Document()
    config = load_template(TEMPLATES_DIR / "thesis.yaml")
    style = register_heading_style(document, "Heading 1", config.headings.level_1)
    assert style.paragraph_format.space_before.pt == 24.0
    assert style.paragraph_format.space_after.pt == 18.0


def test_register_heading_style_supports_right_alignment():
    document = Document()
    config = load_template(TEMPLATES_DIR / "thesis.yaml")
    right_aligned = replace(config.headings.level_1, align="right")

    style = register_heading_style(document, "Heading 1", right_aligned)

    assert style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT
