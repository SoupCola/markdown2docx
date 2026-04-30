from __future__ import annotations

from unittest.mock import MagicMock

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


def test_latex_to_omml_simple():
    """Simple LaTeX x+y converts to OMML XML element."""
    from docx_formatter.formulas import latex_to_omml_element

    el = latex_to_omml_element("x+y")
    assert el is not None
    omml_str = etree.tostring(el, encoding="unicode")
    assert "m:oMath" in omml_str


def test_latex_to_omml_fraction():
    """LaTeX fraction converts correctly."""
    from docx_formatter.formulas import latex_to_omml_element

    el = latex_to_omml_element(r"\frac{1}{2}")
    assert el is not None
    omml_str = etree.tostring(el, encoding="unicode")
    assert "m:f" in omml_str


def test_latex_to_omml_superscript():
    """LaTeX superscript converts correctly."""
    from docx_formatter.formulas import latex_to_omml_element

    el = latex_to_omml_element("x^2")
    assert el is not None
    omml_str = etree.tostring(el, encoding="unicode")
    assert "m:sSup" in omml_str


def test_set_formula_font():
    """Font is set to Times New Roman on all m:r elements."""
    from docx_formatter.formulas import latex_to_omml_element, set_formula_font

    el = latex_to_omml_element("x+y")
    set_formula_font(el, "Times New Roman")
    omml_str = etree.tostring(el, encoding="unicode")
    assert "Times New Roman" in omml_str


def test_is_formula_paragraph_false():
    """Regular paragraph is not a formula paragraph."""
    from docx_formatter.formulas import is_formula_paragraph

    doc = Document()
    p = doc.add_paragraph("Hello world")
    assert is_formula_paragraph(p) is False


def test_is_formula_paragraph_true():
    """Paragraph with inserted formula is detected."""
    from docx_formatter.formulas import insert_formula_paragraph, is_formula_paragraph

    doc = Document()
    p = insert_formula_paragraph(doc, "x+y", "Times New Roman")
    assert is_formula_paragraph(p) is True


def test_detect_formula_paragraphs():
    """Detects formula paragraphs among regular paragraphs."""
    from docx_formatter.formulas import detect_formula_paragraphs, insert_formula_paragraph

    doc = Document()
    doc.add_paragraph("Regular text 1")
    insert_formula_paragraph(doc, "x+y", "Times New Roman")
    doc.add_paragraph("Regular text 2")
    insert_formula_paragraph(doc, r"\frac{a}{b}", "Times New Roman")

    indices = detect_formula_paragraphs(doc)
    assert indices == [1, 3]


def _make_formula_config():
    """Create a mock FormulaConfig for testing."""
    config = MagicMock()
    config.font_cn = "宋体"
    config.font_en = "Times New Roman"
    config.size = "小四"
    config.align = "center"
    config.space_before = "6pt"
    config.space_after = "6pt"
    config.numbering = "chapter"
    return config


def test_apply_formula_formatting_no_center_align():
    """Formula paragraphs do NOT get center alignment (numbering uses tab stops)."""
    from docx_formatter.formulas import apply_formula_formatting, insert_formula_paragraph

    doc = Document()
    doc.add_paragraph("Regular text")
    insert_formula_paragraph(doc, "x+y", "Times New Roman")

    config = _make_formula_config()
    apply_formula_formatting(doc, config)

    # alignment should NOT be set to center (numbering handles it via tabs)
    assert doc.paragraphs[1].paragraph_format.alignment is None


def test_apply_formula_formatting_no_indent():
    """Formula paragraphs have no first-line indent."""
    from docx_formatter.formulas import apply_formula_formatting, insert_formula_paragraph

    doc = Document()
    doc.add_paragraph("Regular text")
    p = insert_formula_paragraph(doc, "x+y", "Times New Roman")

    config = _make_formula_config()
    apply_formula_formatting(doc, config)

    indent = p.paragraph_format.first_line_indent
    assert indent is None or indent == Pt(0)


def test_apply_formula_formatting_skips_regular():
    """Regular paragraphs are not affected by formula formatting."""
    from docx_formatter.formulas import apply_formula_formatting

    doc = Document()
    doc.add_paragraph("Regular text")

    config = _make_formula_config()
    apply_formula_formatting(doc, config)

    assert doc.paragraphs[0].paragraph_format.alignment is None


def test_apply_formula_numbering_chapter():
    """Chapter-based numbering: (1-1), (1-2), (2-1)... with tab stops."""
    from docx_formatter.formulas import apply_formula_numbering, insert_formula_paragraph

    doc = Document()
    # Chapter 1
    doc.add_heading("Chapter 1", level=1)
    insert_formula_paragraph(doc, "x+y", "Times New Roman")  # (1-1)
    insert_formula_paragraph(doc, "a+b", "Times New Roman")  # (1-2)
    # Chapter 2
    doc.add_heading("Chapter 2", level=1)
    insert_formula_paragraph(doc, r"\frac{1}{2}", "Times New Roman")  # (2-1)

    config = _make_formula_config()
    config.numbering = "chapter"
    apply_formula_numbering(doc, config)

    # Formula paragraphs are at indices 1, 2, 4
    p1 = doc.paragraphs[1]
    p2 = doc.paragraphs[2]
    p4 = doc.paragraphs[4]

    texts = []
    for p in [p1, p2, p4]:
        for run in p.runs:
            if run.text.strip():
                texts.append(run.text.strip())
    assert "(1-1)" in texts
    assert "(1-2)" in texts
    assert "(2-1)" in texts

    # Verify tab stops are set on formula paragraphs
    for p in [p1, p2, p4]:
        pPr = p._p.find(qn("w:pPr"))
        tabs = pPr.find(qn("w:tabs")) if pPr is not None else None
        assert tabs is not None
        tab_elements = tabs.findall(qn("w:tab"))
        assert len(tab_elements) == 2
        assert tab_elements[0].get(qn("w:val")) == "center"
        assert tab_elements[1].get(qn("w:val")) == "right"


def test_apply_formula_numbering_continuous():
    """Continuous numbering: (1), (2), (3)..."""
    from docx_formatter.formulas import apply_formula_numbering, insert_formula_paragraph

    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    insert_formula_paragraph(doc, "x+y", "Times New Roman")
    doc.add_heading("Chapter 2", level=1)
    insert_formula_paragraph(doc, "a+b", "Times New Roman")

    config = _make_formula_config()
    config.numbering = "continuous"
    apply_formula_numbering(doc, config)

    p1 = doc.paragraphs[1]
    p2 = doc.paragraphs[3]

    texts = []
    for p in [p1, p2]:
        for run in p.runs:
            if run.text.strip():
                texts.append(run.text.strip())
    assert "(1)" in texts
    assert "(2)" in texts


def test_apply_formula_numbering_registers_labels():
    from docx_formatter.formulas import apply_formula_numbering, insert_formula_paragraph
    from docx_formatter.references import LabelRegistry

    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    insert_formula_paragraph(doc, "x+y", "Times New Roman")
    insert_formula_paragraph(doc, "a+b", "Times New Roman")
    doc.add_heading("Chapter 2", level=1)
    insert_formula_paragraph(doc, r"\frac{1}{2}", "Times New Roman")

    config = _make_formula_config()
    config.numbering = "chapter"

    registry = LabelRegistry()
    formula_items = [
        (1, {"key": "formula:add"}),
        (2, {"key": "formula:sum"}),
        (4, {"key": "formula:frac"}),
    ]

    apply_formula_numbering(doc, config, formula_items=formula_items, registry=registry)

    e1 = registry.lookup("formula:add")
    assert e1 is not None
    assert e1.label_text == "(1-1)"

    e2 = registry.lookup("formula:sum")
    assert e2 is not None
    assert e2.label_text == "(1-2)"

    e3 = registry.lookup("formula:frac")
    assert e3 is not None
    assert e3.label_text == "(2-1)"


def test_apply_formula_numbering_adds_bookmarks():
    from docx_formatter.formulas import apply_formula_numbering, insert_formula_paragraph
    from docx_formatter.references import LabelRegistry

    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    insert_formula_paragraph(doc, "x+y", "Times New Roman")

    config = _make_formula_config()
    config.numbering = "chapter"

    registry = LabelRegistry()
    formula_items = [(1, {"key": "formula:add"})]

    apply_formula_numbering(doc, config, formula_items=formula_items, registry=registry)

    xml = doc.paragraphs[1]._p.xml
    assert "_formula_add" in xml
